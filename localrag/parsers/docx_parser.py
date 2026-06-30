from pathlib import Path
from docx import Document

def parse_docx(path: Path) -> str:
    """Extrae texto de un archivo .docx.
    Si falla, devuelve una cadena vacía y muestra un warning.
    """
    try:
        doc = Document(path)
    except Exception as e:
        print(f"[ADVERTENCIA] No se pudo leer {path}: {e}")
        return ""   # Devuelve string vacío para no romper el flujo

    texto = ""
    for parrafo in doc.paragraphs:
        if parrafo.text.strip():
            texto += parrafo.text + "\n"
    for tabla in doc.tables:
        for fila in tabla.rows:
            texto += " | ".join(celda.text.strip() for celda in fila.cells) + "\n"
    return texto