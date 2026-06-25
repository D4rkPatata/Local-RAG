from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = "/sessions/youthful-brave-goldberg/mnt/14. Local-RAG/corpus/synthetic_docs"

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def header_footer(doc, title):
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = f"Nexus Soluciones S.A.C. — {title}"
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer
    footer.paragraphs[0].text = "Documento Interno — Uso Exclusivo de Colaboradores Nexus Soluciones S.A.C."
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# D06 — Código de Conducta y Ética Empresarial
def gen_D06():
    doc = Document()
    header_footer(doc, "Código de Conducta y Ética Empresarial")
    add_heading(doc, "Código de Conducta y Ética Empresarial", 1)
    doc.add_paragraph("Código: LEGAL-POL-001 | Versión: 2.3 | Fecha: Enero 2026 | Propietario: Gerencia Legal y Compliance")

    add_heading(doc, "1. Propósito", 2)
    add_paragraph(doc, "El presente Código establece los estándares de conducta y valores éticos que todos los colaboradores de Nexus Soluciones S.A.C. deben cumplir en el ejercicio de sus funciones, tanto dentro como fuera de las instalaciones de la empresa.")

    add_heading(doc, "2. Vestimenta y Presentación Personal", 2)
    add_paragraph(doc, "Lunes a jueves: Smart casual. Viernes: Casual. En visitas a clientes o eventos corporativos: Formal o smart formal. El uso de ropa con mensajes ofensivos, político-partidarios o inapropiados está prohibido en todo momento.")

    add_heading(doc, "3. Conflicto de Intereses", 2)
    add_paragraph(doc, "Todo colaborador debe presentar una declaración anual de conflictos de interés al área Legal y Compliance antes del 31 de enero de cada año. Se considera conflicto de interés cualquier situación donde los intereses personales, familiares o financieros del colaborador puedan influir (o aparentar influir) en sus decisiones laborales. Ante cualquier duda, consultar con Mónica Salinas Bustamante: m.salinas@nexussoluciones.pe / Ext. 106.")

    add_heading(doc, "4. Política Anticorrupción y Regalos", 2)
    add_paragraph(doc, "Nexus Soluciones S.A.C. tiene una política de tolerancia cero ante la corrupción. Está absolutamente prohibido:")
    prohib = [
        "Ofrecer, prometer o aceptar pagos, regalos o beneficios para obtener o mantener un contrato.",
        "Aceptar regalos de proveedores o clientes con valor superior a S/. 50 en total por año calendario.",
        "Invitaciones a eventos (deportivos, sociales, entretenimiento) de valor superior a S/. 80 sin aprobación del Gerente General.",
        "Realizar pagos a terceros con el propósito de facilitar trámites o decisiones que beneficien a la empresa.",
    ]
    for p in prohib:
        doc.add_paragraph(f"• {p}", style="List Bullet")

    add_heading(doc, "5. Uso de Redes Sociales", 2)
    add_paragraph(doc, "Los colaboradores no deben publicar en redes sociales personales información sobre proyectos en curso, nombres de clientes, datos internos de la empresa, resultados financieros no publicados o fotografías de instalaciones restringidas. El uso de redes sociales en horario laboral debe ser moderado y no interferir con el desempeño.")

    add_heading(doc, "6. Trato entre Colaboradores", 2)
    add_paragraph(doc, "Nexus promueve un ambiente laboral de respeto y colaboración. Están prohibidos el acoso laboral (mobbing), el acoso sexual, la discriminación por cualquier condición personal, y toda forma de violencia física o verbal. Las situaciones de conflicto entre colaboradores deben reportarse a RR.HH. (s.torres@nexussoluciones.pe).")

    add_heading(doc, "7. Proceso Disciplinario", 2)
    add_paragraph(doc, "El incumplimiento del presente Código activa el siguiente proceso disciplinario progresivo:")
    niveles = [
        ("Nivel 1", "Amonestación verbal", "Falta leve por primera vez"),
        ("Nivel 2", "Amonestación escrita", "Reincidencia en falta leve o falta moderada"),
        ("Nivel 3", "Suspensión sin goce de haber (1 a 5 días)", "Falta grave o reincidencia en falta moderada"),
        ("Nivel 4", "Despido por causa justificada", "Falta muy grave o reincidencia tras suspensión"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["Nivel", "Medida", "Aplicable cuando"]):
        t.rows[0].cells[i].text = h
    for nivel, medida, cuando in niveles:
        row = t.add_row()
        row.cells[0].text = nivel
        row.cells[1].text = medida
        row.cells[2].text = cuando

    add_heading(doc, "8. Canal de Denuncias (Línea Ética)", 2)
    add_paragraph(doc, "Nexus dispone de un canal de denuncias anónimo para reportar incumplimientos éticos, actos de corrupción o cualquier conducta contraria a este Código:")
    add_paragraph(doc, "• Correo confidencial: denuncias@nexussoluciones.pe\n• Línea ética (llamada gratuita): 0800-00-555\n\nTodas las denuncias son tratadas con estricta confidencialidad y no podrán tomarse represalias contra quien denuncia de buena fe.")

    doc.save(f"{OUTPUT_DIR}/D06_Codigo_Conducta_Etica.docx")
    print("D06 OK")

# D07 — Plan de Capacitación y Desarrollo Profesional
def gen_D07():
    doc = Document()
    header_footer(doc, "Plan de Capacitación y Desarrollo Profesional")
    add_heading(doc, "Plan de Capacitación y Desarrollo Profesional", 1)
    doc.add_paragraph("Código: RRHH-POL-005 | Versión: 2.0 | Fecha: Enero 2026 | Propietario: Gerencia de RR.HH.")

    add_heading(doc, "1. Objetivo", 2)
    add_paragraph(doc, "Garantizar el desarrollo continuo de las competencias técnicas y blandas de todos los colaboradores de Nexus Soluciones S.A.C., alineando el crecimiento individual con los objetivos estratégicos de la empresa.")

    add_heading(doc, "2. Presupuesto de Capacitación", 2)
    add_paragraph(doc, "Cada colaborador cuenta con un presupuesto anual de S/. 1,500 para capacitación externa. Este monto se gestiona a través del área de RR.HH. y no es acumulable entre años. Las solicitudes que superen los S/. 800 requieren aprobación adicional del CFO, Patricia Quispe Huamán (p.quispe@nexussoluciones.pe).")

    add_heading(doc, "3. Nexus Academy — Plataforma E-learning Interna", 2)
    add_paragraph(doc, "La plataforma de aprendizaje en línea de Nexus está disponible en intranet.nexussoluciones.pe/academy. Incluye cursos técnicos, de liderazgo, seguridad de la información y normativas internas. El acceso es gratuito para todos los colaboradores en planilla y está disponible las 24 horas.")

    add_heading(doc, "4. Certificaciones Pagadas por la Empresa", 2)
    add_paragraph(doc, "La empresa cubre el 100% del costo de las siguientes certificaciones estratégicas, sujeto a disponibilidad presupuestal y aprobación del jefe directo:")
    certs = ["AWS (Cloud Practitioner, Solutions Architect)", "Microsoft Azure (Fundamentals, Administrator)", "Certified Scrum Master (CSM)", "Project Management Professional (PMP)", "ISO 27001 Lead Implementer"]
    for c in certs:
        doc.add_paragraph(f"• {c}", style="List Bullet")
    add_paragraph(doc, "El colaborador que recibe financiamiento de certificación asume el compromiso de permanecer en la empresa por un mínimo de 12 meses a partir de la fecha de obtención del certificado.")

    add_heading(doc, "5. Clases de Inglés Corporativas", 2)
    add_paragraph(doc, "Nexus ofrece clases de inglés gratuitas para todos sus colaboradores en los siguientes horarios:\n• Martes y jueves: 7:00 – 8:00 am (modalidad virtual)\nLas clases están a cargo de un proveedor externo. La inscripción se realiza en intranet.nexussoluciones.pe/academy/ingles. Se requiere asistencia mínima del 80% para mantener la matrícula.")

    add_heading(doc, "6. Proceso de Solicitud de Capacitación Externa", 2)
    pasos = [
        "El colaborador identifica la capacitación y presenta una solicitud en intranet.nexussoluciones.pe/capacitacion, indicando nombre del curso, proveedor, costo, fechas y justificación.",
        "El jefe directo evalúa la pertinencia y aprueba o rechaza en 5 días hábiles.",
        "Si el costo supera S/. 800, la solicitud pasa automáticamente al CFO para aprobación.",
        "RR.HH. coordina el pago y el registro en el expediente del colaborador.",
        "Al finalizar, el colaborador debe compartir un resumen o presentación de lo aprendido con su equipo (30 minutos máximo).",
    ]
    for i, p in enumerate(pasos, 1):
        doc.add_paragraph(f"{i}. {p}")

    add_heading(doc, "7. Plan de Desarrollo Individual (PDI)", 2)
    add_paragraph(doc, "Al cierre de cada Evaluación Year-End, el jefe directo y el colaborador definen en conjunto un Plan de Desarrollo Individual para el año siguiente. El PDI incluye objetivos de desarrollo, competencias a fortalecer, capacitaciones planificadas y un responsable de seguimiento. El PDI se registra en intranet.nexussoluciones.pe/evaluacion y es revisado en la evaluación Mid-Year.")

    doc.save(f"{OUTPUT_DIR}/D07_Plan_Capacitacion.docx")
    print("D07 OK")

# D08 — Reglamento Interno de SST
def gen_D08():
    doc = Document()
    header_footer(doc, "Reglamento Interno de Seguridad y Salud en el Trabajo")
    add_heading(doc, "Reglamento Interno de Seguridad y Salud en el Trabajo", 1)
    doc.add_paragraph("Código: SST-REG-001 | Versión: 3.0 | Fecha: Enero 2026 | Propietario: Comité de SST")

    add_heading(doc, "1. Fundamento Legal", 2)
    add_paragraph(doc, "El presente Reglamento se rige por la Ley N.° 29783 — Ley de Seguridad y Salud en el Trabajo y su Reglamento aprobado por D.S. N.° 005-2012-TR, y sus modificatorias.")

    add_heading(doc, "2. Comité de Seguridad y Salud en el Trabajo", 2)
    add_paragraph(doc, "El Comité de SST está conformado por representantes de la empresa y de los colaboradores. El representante de la empresa es Marco Huanca Soto (m.huanca@nexussoluciones.pe, Ext. 301). El representante de los colaboradores es elegido anualmente por votación de todos los colaboradores en planilla.")

    add_heading(doc, "3. Obligaciones de los Colaboradores", 2)
    oblig = [
        "Cumplir las normas, reglamentos y procedimientos del Sistema de Gestión de SST.",
        "Usar correctamente los equipos de protección personal asignados.",
        "No manipular ni deshabilitar dispositivos de seguridad (extintores, alarmas, AED).",
        "Reportar inmediatamente cualquier accidente, incidente o condición insegura al Comité de SST.",
        "Participar en los simulacros de evacuación programados (mínimo 2 al año).",
        "Pasar los exámenes ocupacionales anuales coordinados por RR.HH. en febrero de cada año.",
    ]
    for o in oblig:
        doc.add_paragraph(f"• {o}", style="List Bullet")

    add_heading(doc, "4. Infraestructura de Emergencia", 2)
    add_paragraph(doc, "Botiquines de primeros auxilios:", bold=True)
    bot = ["Recepción San Isidro — Piso 8", "Piso 9 San Isidro (área TI)", "Recepción Miraflores — Piso 3"]
    for b in bot:
        doc.add_paragraph(f"• {b}", style="List Bullet")
    add_paragraph(doc, "Desfibrilador (AED):", bold=True)
    add_paragraph(doc, "Ubicado en Recepción San Isidro, Piso 8. Solo personal capacitado debe operarlo.")
    add_paragraph(doc, "Salidas de emergencia:", bold=True)
    add_paragraph(doc, "Las escaleras de emergencia son accesibles por el pasillo central de cada piso. Están señalizadas con luces de emergencia. Los ascensores NO deben usarse durante una evacuación.")
    add_paragraph(doc, "Punto de reunión:", bold=True)
    add_paragraph(doc, "Estacionamiento B1, zona señalizada con la letra 'E'. Todos los colaboradores deben dirigirse ahí en caso de evacuación.")

    add_heading(doc, "5. Simulacros de Evacuación", 2)
    add_paragraph(doc, "Se realizan 2 simulacros anuales de evacuación. El primero en el primer semestre (marzo–abril) y el segundo en el segundo semestre (septiembre–octubre). La participación es obligatoria para todos los colaboradores presentes. El tiempo objetivo de evacuación total es de 4 minutos.")

    add_heading(doc, "6. Exámenes Médicos Ocupacionales", 2)
    add_paragraph(doc, "Todos los colaboradores realizan exámenes médicos ocupacionales anuales coordinados por RR.HH. en el mes de febrero. Los resultados son confidenciales y solo son accesibles por el médico ocupacional y el propio colaborador.")

    add_heading(doc, "7. Ergonomía en el Puesto de Trabajo", 2)
    add_paragraph(doc, "TI provee sillas ergonómicas y pantallas a la altura adecuada en todas las estaciones de trabajo. Para trabajadores remotos con más de 6 meses en la empresa, se puede solicitar una evaluación ergonómica del puesto en casa, coordinando con el área de SST a través de m.huanca@nexussoluciones.pe.")

    add_heading(doc, "8. Reporte de Accidentes e Incidentes", 2)
    add_paragraph(doc, "Todo accidente o incidente, por leve que sea, debe reportarse el mismo día al Comité de SST (m.huanca@nexussoluciones.pe). RR.HH. emite el aviso de accidente de trabajo ante ESSALUD en un plazo máximo de 5 días hábiles desde ocurrido el evento.")

    doc.save(f"{OUTPUT_DIR}/D08_Reglamento_SST.docx")
    print("D08 OK")

# D09 — Política de Seguridad de la Información y Uso de Equipos TI
def gen_D09():
    doc = Document()
    header_footer(doc, "Política de Seguridad de la Información y Uso de Equipos TI")
    add_heading(doc, "Política de Seguridad de la Información y Uso de Equipos TI", 1)
    doc.add_paragraph("Código: TI-POL-001 | Versión: 4.1 | Fecha: Enero 2026 | Propietario: CISO — Diana Reyes Castañeda")

    add_heading(doc, "1. Objetivo y Alcance", 2)
    add_paragraph(doc, "Esta política establece las normas de uso seguro de los equipos, sistemas y datos de Nexus Soluciones S.A.C., alineadas con la certificación ISO 27001:2022. Aplica a todos los colaboradores, locadores de servicios y terceros con acceso a los sistemas de la empresa.")

    add_heading(doc, "2. Equipos Corporativos Asignados", 2)
    add_paragraph(doc, "Modelo estándar: Lenovo ThinkPad E14 Gen 4 o Dell Latitude 5540. Sistema operativo: Windows 11 Pro con imagen corporativa. Todos los equipos cuentan con BitLocker habilitado (cifrado de disco completo) y CrowdStrike Falcon como solución de seguridad endpoint. No se debe desinstalar ni deshabilitar ninguno de estos componentes.")

    add_heading(doc, "3. Política de Contraseñas", 2)
    add_paragraph(doc, "Todos los accesos a sistemas corporativos deben cumplir los siguientes requisitos de contraseña:")
    req_pass = [
        "Longitud mínima: 12 caracteres",
        "Debe incluir: mayúsculas, minúsculas, números y caracteres especiales (!@#$%^&*)",
        "Caducidad: cada 90 días",
        "No reutilizar las últimas 6 contraseñas",
        "MFA (Autenticación Multifactor) obligatorio mediante Microsoft Authenticator",
    ]
    for r in req_pass:
        doc.add_paragraph(f"• {r}", style="List Bullet")

    add_heading(doc, "4. VPN Corporativa", 2)
    add_paragraph(doc, "Para acceder a los sistemas internos desde fuera de la oficina es obligatorio usar Cisco AnyConnect. La solicitud de acceso VPN se realiza a helpdesk@nexussoluciones.pe con aprobación del jefe directo, a través de intranet.nexussoluciones.pe/solicitudes-ti.")

    add_heading(doc, "5. Herramientas Autorizadas y Prohibidas", 2)
    add_paragraph(doc, "Herramientas autorizadas:", bold=True)
    auth = [
        "Comunicación: Microsoft Teams, Outlook",
        "Gestión de proyectos: Jira, Confluence",
        "Repositorios de código: GitLab interno (gitlab.nexussoluciones.pe), GitHub (solo repos autorizados)",
        "Almacenamiento en nube: SharePoint corporativo, OneDrive personal corporativo",
        "Videoconferencias: Microsoft Teams, Zoom con cuenta corporativa",
    ]
    for a in auth:
        doc.add_paragraph(f"• {a}", style="List Bullet")
    add_paragraph(doc, "Herramientas y usos prohibidos:", bold=True)
    prohib = [
        "Software de torrents o descarga P2P",
        "Software de minería de criptomonedas",
        "VPNs no autorizadas por TI",
        "Software sin licencia o crackeado",
        "Google Drive personal, Dropbox, WeTransfer para información sensible",
        "Herramientas de IA generativa con datos de clientes (requiere aprobación del CISO para proyectos internos)",
    ]
    for p in prohib:
        doc.add_paragraph(f"• {p}", style="List Bullet")

    add_heading(doc, "6. Clasificación de la Información", 2)
    clasificacion = [
        ("Confidencial", "Datos de clientes, contratos, info financiera, código fuente de proyectos", "Solo acceso con necesidad justificada, no enviar sin cifrar, no almacenar en dispositivos personales"),
        ("Interno", "Políticas, procedimientos, comunicaciones internas", "Solo para colaboradores de Nexus, no compartir sin aprobación"),
        ("Público", "Material de marketing, web corporativa", "Sin restricciones"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["Nivel", "Ejemplos", "Manejo requerido"]):
        t.rows[0].cells[i].text = h
    for row_data in clasificacion:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    add_heading(doc, "7. Contacto TI", 2)
    add_paragraph(doc, "Mesa de Ayuda (Helpdesk): helpdesk@nexussoluciones.pe | Ext. 300\nHorario: Lun–Vie 8:00–19:00 | Sáb 9:00–13:00\nCISO: Diana Reyes Castañeda — d.reyes@nexussoluciones.pe | Ext. 302\nLíder de Infraestructura: Marco Huanca Soto — m.huanca@nexussoluciones.pe | Ext. 301")

    doc.save(f"{OUTPUT_DIR}/D09_Politica_Seguridad_TI.docx")
    print("D09 OK")

# D10 — Procedimiento de Gestión de Incidentes de TI
def gen_D10():
    doc = Document()
    header_footer(doc, "Procedimiento de Gestión de Incidentes de TI")
    add_heading(doc, "Procedimiento de Gestión de Incidentes de TI", 1)
    doc.add_paragraph("Código: TI-PROC-001 | Versión: 3.5 | Fecha: Enero 2026 | Propietario: CISO — Diana Reyes Castañeda")

    add_heading(doc, "1. Definición de Incidente", 2)
    add_paragraph(doc, "Un incidente de TI es cualquier evento no planificado que afecte o pueda afectar la disponibilidad, confidencialidad o integridad de los sistemas, datos o infraestructura de Nexus Soluciones S.A.C. o de sus clientes.")

    add_heading(doc, "2. Clasificación de Incidentes por Severidad", 2)
    niveles = [
        ("P1 — Crítico", "Sistema productivo caído, brecha de seguridad confirmada, pérdida de datos de cliente", "15 minutos", "4 horas", "CISO, CTO, Gerente General"),
        ("P2 — Alto", "Degradación significativa de servicio, sospecha de intrusión, acceso no autorizado", "30 minutos", "8 horas", "CISO, CTO"),
        ("P3 — Medio", "Problema de acceso individual, fallo de equipo, error en app interna", "2 horas", "24 horas", "Líder de Infraestructura"),
        ("P4 — Bajo", "Consultas, solicitudes de acceso, instalación de software", "8 horas", "48 horas", "Mesa de Ayuda"),
    ]
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    for i, h in enumerate(["Nivel", "Descripción", "Tiempo de Respuesta", "Tiempo de Resolución", "Notificar a"]):
        t.rows[0].cells[i].text = h
    for row_data in niveles:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    add_heading(doc, "3. Cómo Reportar un Incidente", 2)
    add_paragraph(doc, "Todo colaborador que identifique o sospeche de un incidente debe reportarlo inmediatamente por los siguientes canales:")
    canales = [
        "Correo: helpdesk@nexussoluciones.pe",
        "Teléfono: Ext. 300 (Helpdesk) — Lun–Vie 8:00–19:00, Sáb 9:00–13:00",
        "Portal de tickets: jira.nexussoluciones.pe/servicedesk",
        "Para incidentes P1 fuera de horario: llamar directamente al CISO (d.reyes@nexussoluciones.pe) o al CTO (a.flores@nexussoluciones.pe)",
    ]
    for c in canales:
        doc.add_paragraph(f"• {c}", style="List Bullet")

    add_heading(doc, "4. Proceso de Atención", 2)
    pasos = [
        ("Registro", "El colaborador reporta el incidente. El Helpdesk crea un ticket en Jira Service Management con severidad inicial asignada."),
        ("Clasificación", "TI valida la severidad real en un tiempo máximo igual al tiempo de respuesta del nivel asignado."),
        ("Escalamiento", "Si el incidente es P1 o P2, se notifica de inmediato a los responsables definidos en la tabla de severidad."),
        ("Resolución", "El equipo de TI trabaja en la solución dentro del tiempo de resolución comprometido. Se mantiene al usuario informado cada hora en incidentes P1."),
        ("Cierre", "TI confirma la resolución con el usuario afectado y cierra el ticket en Jira."),
        ("Post-incidente", "Para P1 y P2, TI emite un informe post-incidente (RCA) en un plazo máximo de 5 días hábiles."),
    ]
    for paso, desc in pasos:
        add_paragraph(doc, paso, bold=True)
        add_paragraph(doc, desc)

    add_heading(doc, "5. Incidentes de Seguridad — Protocolo Especial", 2)
    add_paragraph(doc, "Ante la sospecha o confirmación de una brecha de seguridad que involucre datos de clientes:")
    proto = [
        "El colaborador NO debe intentar resolver el incidente por su cuenta.",
        "Debe reportar inmediatamente al CISO: d.reyes@nexussoluciones.pe / Ext. 302.",
        "El CISO notifica al Oficial de Protección de Datos (DPO): dpo@nexussoluciones.pe.",
        "Si se confirma exposición de datos personales de clientes, la notificación al cliente y a INDECOPI debe realizarse en un plazo máximo de 72 horas.",
    ]
    for p in proto:
        doc.add_paragraph(f"• {p}", style="List Bullet")

    doc.save(f"{OUTPUT_DIR}/D10_Gestion_Incidentes_TI.docx")
    print("D10 OK")

gen_D06()
gen_D07()
gen_D08()
gen_D09()
gen_D10()
print("Batch 2 (D06-D10) completado.")
