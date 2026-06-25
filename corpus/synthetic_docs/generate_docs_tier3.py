"""
generate_docs_tier3.py — Generador de documentos Tier-3 (D16–D23) para Nexus Soluciones S.A.C.

Documentos confidenciales / estratégicos que justifican el carácter local/offline del chatbot.
Toda la data deriva de corpus/ground_truth/nexus_ground_truth_tier3.json.

Mantiene el estilo y patrón de generate_docs_2.py / generate_docs_3.py.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Por defecto guarda al lado del script — funciona en cualquier OS / sesión.
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ----- helpers (mismo patrón que generate_docs_*.py) -----

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def header_footer(doc, title, classification="CONFIDENCIAL — Uso Interno Restringido"):
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = f"Nexus Soluciones S.A.C. — {title}"
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer
    footer.paragraphs[0].text = (
        f"{classification} — Prohibida su reproducción o divulgación externa sin autorización."
    )
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    return t


# ===================================================================
# D16 — Modelo de Pricing y Tarifario Profesional por Seniority
# ===================================================================
def gen_D16():
    doc = Document()
    header_footer(doc, "Modelo de Pricing y Tarifario")
    add_heading(doc, "Modelo de Pricing y Tarifario Profesional por Seniority", 1)
    doc.add_paragraph(
        "Código: COMERCIAL-POL-001 | Versión: 5.2 | Fecha: Enero 2026 | "
        "Propietario: Gerencia Comercial — Carlos Amat y León Ríos | "
        "Clasificación: CONFIDENCIAL"
    )

    add_heading(doc, "1. Objetivo y Alcance", 2)
    add_paragraph(
        doc,
        "Establecer las tarifas profesionales por rol y seniority aplicables a propuestas "
        "comerciales de Nexus Soluciones S.A.C., así como los márgenes mínimos, política de "
        "descuentos, premium por industria y recargos por horarios no estándar. Aplica a todas "
        "las modalidades de contratación: tiempo y materiales (T&M), precio fijo (Fixed Price) "
        "y aumento de personal (Staff Augmentation).",
    )

    add_heading(doc, "2. Tarifas Hora Estándar por Rol", 2)
    add_paragraph(doc, "Tarifas expresadas en Soles peruanos (S/.) por hora trabajada:")
    add_table(doc, ["Rol", "Tarifa Hora (S/.)"], [
        ("Desarrollador Junior", 45),
        ("Desarrollador Semi-Senior", 75),
        ("Desarrollador Senior", 110),
        ("Tech Lead", 145),
        ("Software Architect", 180),
        ("QA Junior", 40),
        ("QA Senior", 95),
        ("Project Manager", 130),
        ("Consultor Senior", 160),
        ("Consultor Principal", 210),
    ])

    add_heading(doc, "3. Costo Total Interno (CTC) por Rol", 2)
    add_paragraph(
        doc,
        "Información estrictamente confidencial. El CTC es la base interna para el cálculo de "
        "margen bruto en propuestas. Su divulgación a clientes o terceros constituye falta grave.",
    )
    add_table(doc, ["Rol", "CTC Hora (S/.)"], [
        ("Desarrollador Junior", 28),
        ("Desarrollador Semi-Senior", 46),
        ("Desarrollador Senior", 68),
        ("Tech Lead", 92),
        ("Software Architect", 115),
        ("Consultor Senior", 102),
        ("Consultor Principal", 140),
    ])

    add_heading(doc, "4. Margen Mínimo por Modalidad", 2)
    add_table(doc, ["Modalidad", "Margen mínimo aceptable"], [
        ("Tiempo y Materiales (T&M)", "38%"),
        ("Precio Fijo (Fixed Price)", "42%"),
        ("Staff Augmentation", "28%"),
    ])
    add_paragraph(
        doc,
        "Toda propuesta que no alcance el margen mínimo requiere aprobación expresa del CEO "
        "y debe sustentarse en valor estratégico (entrada a nueva cuenta, cierre de competidor, etc.).",
    )

    add_heading(doc, "5. Política de Descuentos Autorizados", 2)
    add_table(doc, ["Rango de descuento", "Aprobador requerido"], [
        ("0% – 5%", "Gerente Comercial"),
        ("5.01% – 10%", "Gerente General (CEO)"),
        ("Mayor a 10%", "Directorio"),
    ])

    add_heading(doc, "6. Premium por Industria Regulada", 2)
    add_paragraph(
        doc,
        "Los sectores con alta regulación de datos y compliance reciben un recargo sobre la tarifa "
        "estándar, en reconocimiento del esfuerzo adicional de cumplimiento normativo y seguridad:",
    )
    add_table(doc, ["Industria", "Premium aplicable"], [
        ("Banca y Finanzas", "+15%"),
        ("Salud", "+12%"),
        ("Seguros", "+10%"),
        ("Retail", "+5%"),
    ])

    add_heading(doc, "7. Recargos por Horarios No Estándar", 2)
    add_table(doc, ["Tipo de jornada", "Recargo sobre tarifa estándar"], [
        ("Fuera de horario (lunes–viernes 19:00–07:00)", "+50%"),
        ("Fines de semana (sábado y domingo)", "+100%"),
        ("Emergencia inmediata (atención <2 h)", "+150%"),
    ])

    add_heading(doc, "8. Ciclo de Revisión de Tarifas", 2)
    add_paragraph(
        doc,
        "La revisión ordinaria de tarifas se realiza anualmente cada enero. Se gatilla una "
        "revisión extraordinaria si el IPC acumulado supera 5% o si el tipo de cambio USD/PEN "
        "supera 4.20. La aprobación de la nueva tabla corresponde al CEO previa propuesta de la "
        "Gerencia Comercial.",
    )

    doc.save(f"{OUTPUT_DIR}/D16_Pricing_Tarifario.docx")
    print("D16 OK")


# ===================================================================
# D17 — Playbook de Propuestas Comerciales y RFP
# ===================================================================
def gen_D17():
    doc = Document()
    header_footer(doc, "Playbook de Propuestas y RFP")
    add_heading(doc, "Playbook de Propuestas Comerciales y Respuesta a RFP", 1)
    doc.add_paragraph(
        "Código: COMERCIAL-PROC-001 | Versión: 3.0 | Fecha: Febrero 2026 | "
        "Propietario: Gerencia Comercial | Clasificación: CONFIDENCIAL"
    )

    add_heading(doc, "1. Pipeline Comercial Estandarizado", 2)
    add_paragraph(
        doc,
        "Toda oportunidad comercial debe avanzar por las siguientes etapas, registradas en el CRM "
        "interno (Nexus Pipeline):",
    )
    etapas = ["Lead", "Qualify", "Discovery", "Propuesta", "Negociación", "Cierre", "Handoff a Delivery"]
    for i, e in enumerate(etapas, 1):
        doc.add_paragraph(f"{i}. {e}")

    add_heading(doc, "2. SLA Interno de Respuesta a RFP", 2)
    add_table(doc, ["Monto estimado", "Plazo máximo de respuesta"], [
        ("Menor a S/. 200,000", "5 días hábiles"),
        ("Mayor a S/. 200,000", "10 días hábiles"),
    ])
    add_paragraph(
        doc,
        "El plazo se cuenta desde la recepción formal del RFP. Toda excepción debe ser autorizada "
        "por la Gerencia Comercial.",
    )

    add_heading(doc, "3. Aprobaciones por Monto de Propuesta", 2)
    add_table(doc, ["Monto de la propuesta", "Aprobador final"], [
        ("Hasta S/. 100,000", "Gerente Comercial"),
        ("S/. 100,001 – S/. 500,000", "Gerente General (CEO)"),
        ("Mayor a S/. 500,000", "Directorio"),
    ])

    add_heading(doc, "4. Plantilla Obligatoria de Propuesta", 2)
    secciones = [
        "Resumen ejecutivo",
        "Entendimiento del negocio del cliente",
        "Enfoque y metodología",
        "Equipo propuesto con CV de roles clave",
        "Cronograma e hitos",
        "Inversión y modalidad",
        "Supuestos, exclusiones y dependencias",
    ]
    for s in secciones:
        doc.add_paragraph(f"• {s}", style="List Bullet")

    add_heading(doc, "5. Matriz Go/No-Go", 2)
    add_paragraph(
        doc,
        "Antes de invertir esfuerzo en una propuesta, el Comité Comercial evalúa la oportunidad "
        "contra ocho criterios ponderados:",
    )
    criterios = [
        "Encaje técnico con el stack aprobado (ver D19).",
        "Margen estimado igual o superior al mínimo por modalidad (ver D16).",
        "Riesgo país y solvencia del cliente.",
        "Capacidad disponible en el período propuesto.",
        "Alineación estratégica con sectores foco (Banca, Salud, Retail, Seguros).",
        "Términos contractuales aceptables (ver red flags).",
        "Existencia de referencia o caso previo aplicable.",
        "Probabilidad estimada de cierre igual o mayor a 30%.",
    ]
    for i, c in enumerate(criterios, 1):
        doc.add_paragraph(f"{i}. {c}")

    add_heading(doc, "6. Red Flags Contractuales", 2)
    add_paragraph(
        doc,
        "Las siguientes cláusulas requieren revisión obligatoria de Legal y suelen ser motivo de "
        "rechazo o renegociación:",
    )
    red_flags = [
        "Penalidades superiores al 10% del valor del contrato.",
        "Cesión irrestricta de propiedad intelectual de código o conocimiento.",
        "Jurisdicción extranjera sin cláusula de arbitraje en Lima.",
        "Exclusividad no compensada económicamente.",
        "Indemnización ilimitada por daños.",
    ]
    for r in red_flags:
        doc.add_paragraph(f"• {r}", style="List Bullet")

    add_heading(doc, "7. Lessons Learned — Propuestas Perdidas (Resumen)", 2)
    add_paragraph(
        doc,
        "Casos representativos de propuestas perdidas en los últimos 24 meses y su causa raíz "
        "principal (información disociada del cliente, uso interno de aprendizaje):",
    )
    add_table(doc, ["Sector", "Modalidad", "Causa raíz principal"], [
        ("Banca", "Fixed Price", "Estimación optimista del esfuerzo de integración core bancario."),
        ("Retail", "T&M", "Tarifa hora 18% sobre el promedio del mercado sin diferencial técnico claro."),
        ("Salud", "Fixed Price", "Falta de referencia previa en módulo de historia clínica electrónica."),
    ])

    doc.save(f"{OUTPUT_DIR}/D17_Playbook_RFP.docx")
    print("D17 OK")


# ===================================================================
# D18 — Cartera de Clientes Activos y Maestro de Contratos
# ===================================================================
def gen_D18():
    doc = Document()
    header_footer(doc, "Cartera de Clientes y Contratos",
                  classification="ESTRICTAMENTE CONFIDENCIAL — Acceso por NDA")
    add_heading(doc, "Cartera de Clientes Activos y Maestro de Contratos", 1)
    doc.add_paragraph(
        "Código: COMERCIAL-REG-001 | Versión: 4.1 | Fecha: Marzo 2026 | "
        "Propietario: Gerencia Comercial + CFO | Clasificación: ESTRICTAMENTE CONFIDENCIAL"
    )

    add_heading(doc, "1. Top 10 Clientes Activos 2026", 2)
    add_paragraph(
        doc,
        "Cartera estratégica de Nexus al cierre del primer trimestre 2026. Los montos consignados "
        "corresponden al valor anualizado del contrato vigente.",
    )
    add_table(
        doc,
        ["Cliente", "Sector", "Contrato Anual (S/.)", "Sponsor Nexus", "Inicio", "Renovación"],
        [
            ("Banco Andino del Perú", "Banca y Finanzas", "1,850,000",
             "Carlos Amat y León Ríos", "2022-04-01", "2027-03-31"),
            ("Retail Norte SAC", "Retail", "1,240,000",
             "Patricia Quispe Huamán", "2023-01-15", "2026-12-31"),
            ("Clínica San Borja", "Salud", "980,000",
             "Andrés Flores Castillo", "2023-07-01", "2026-06-30"),
            ("Seguros Pacífico Sur", "Seguros", "760,000",
             "Carlos Amat y León Ríos", "2024-02-01", "2027-01-31"),
            ("MicroFinanzas Inka", "Banca y Finanzas", "620,000",
             "Patricia Quispe Huamán", "2024-09-01", "2026-08-31"),
        ],
    )

    add_heading(doc, "2. Concentración de Ingresos", 2)
    add_paragraph(
        doc,
        "Política de gestión de riesgo de concentración: ningún cliente puede representar más del "
        "25% de los ingresos anuales totales. Actualmente la concentración del cliente top es de "
        "20.1%, dentro del umbral aceptable.",
    )

    add_heading(doc, "3. SLA Contractual por Cliente", 2)
    add_table(doc, ["Cliente", "Uptime comprometido", "Penalidad por incumplimiento"], [
        ("Banco Andino del Perú", "99.7%", "8% del contrato anual por evento"),
        ("Retail Norte SAC", "99.5%", "5% del contrato anual por evento"),
        ("Clínica San Borja", "99.8%", "10% del contrato anual por evento"),
        ("Seguros Pacífico Sur", "99.5%", "6% del contrato anual por evento"),
        ("MicroFinanzas Inka", "99.0%", "4% del contrato anual por evento"),
    ])

    add_heading(doc, "4. Acuerdos de Confidencialidad (NDAs)", 2)
    add_paragraph(
        doc,
        "Nexus mantiene 18 NDAs activos al corte de marzo 2026, incluyendo NDAs mutuos y "
        "unilaterales con clientes, prospectos en negociación y proveedores estratégicos. El "
        "registro maestro de NDAs es administrado por la Gerencia Legal.",
    )

    add_heading(doc, "5. Riesgos Comerciales Abiertos", 2)
    riesgos = [
        "Renovación en riesgo: Clínica San Borja presenta demoras significativas en el proyecto "
        "HIS, lo cual ha generado fricción con el sponsor del cliente.",
        "Cliente moroso: MicroFinanzas Inka registra 45 días de atraso en el pago de la última "
        "factura mensual.",
        "Disputa activa: Retail Norte SAC ha formalizado reclamo por 3 incumplimientos de SLA "
        "durante el Q1-2026, con potencial aplicación de penalidad acumulada.",
    ]
    for r in riesgos:
        doc.add_paragraph(f"• {r}", style="List Bullet")

    add_heading(doc, "6. Plan de Retención de Cuentas Estratégicas", 2)
    add_paragraph(
        doc,
        "Cada cuenta del Top 5 cuenta con un Account Plan trimestral revisado por la Gerencia "
        "Comercial. El plan incluye health score, oportunidades de upsell/cross-sell, riesgos "
        "específicos y acciones de relación ejecutiva. Para detalle por cuenta consultar el "
        "repositorio comercial interno.",
    )

    doc.save(f"{OUTPUT_DIR}/D18_Cartera_Clientes_Contratos.docx")
    print("D18 OK")


# ===================================================================
# D19 — Arquitecturas de Referencia y Estándares Técnicos (ADRs)
# ===================================================================
def gen_D19():
    doc = Document()
    header_footer(doc, "Arquitecturas de Referencia y Estándares Técnicos")
    add_heading(doc, "Arquitecturas de Referencia y Estándares Técnicos (ADRs)", 1)
    doc.add_paragraph(
        "Código: TI-EST-001 | Versión: 2.7 | Fecha: Febrero 2026 | "
        "Propietario: CTO — Andrés Flores Castillo | Clasificación: CONFIDENCIAL"
    )

    add_heading(doc, "1. Stack Tecnológico Aprobado", 2)
    add_paragraph(doc, "Lenguajes:", bold=True)
    add_paragraph(doc, "Python 3.11+, Java 17+, TypeScript 5+, .NET 8.")
    add_paragraph(doc, "Frameworks Backend:", bold=True)
    add_paragraph(doc, "FastAPI, Spring Boot, NestJS, ASP.NET Core.")
    add_paragraph(doc, "Frameworks Frontend:", bold=True)
    add_paragraph(doc, "React 18+, Next.js 14+, Angular 17+.")
    add_paragraph(doc, "Bases de datos:", bold=True)
    add_paragraph(doc, "PostgreSQL 15+ (relacional por defecto), MongoDB 6+, Redis 7+, SQL Server 2022.")
    add_paragraph(doc, "Cloud:", bold=True)
    add_paragraph(
        doc,
        "AWS es el cloud primario aprobado. Azure es secundario y solo se utiliza por requisito "
        "del cliente. GCP queda disponible exclusivamente por excepción autorizada por el CTO.",
    )
    add_paragraph(doc, "Mensajería:", bold=True)
    add_paragraph(doc, "Apache Kafka, AWS SQS, RabbitMQ.")

    add_heading(doc, "2. Patrones Técnicos Obligatorios", 2)
    add_table(doc, ["Aspecto", "Estándar obligatorio"], [
        ("Autenticación", "OAuth 2.0 + OIDC."),
        ("Logging", "Estructurado en JSON con correlation_id."),
        ("Observabilidad", "OpenTelemetry obligatorio en todo servicio productivo."),
        ("Gestión de secretos", "HashiCorp Vault o AWS Secrets Manager. Prohibido en código o "
                                "variables de entorno en claro."),
    ])

    add_heading(doc, "3. Estándares de Seguridad de Código", 2)
    add_paragraph(doc, "Nivel base obligatorio: OWASP ASVS Nivel 2.")
    add_paragraph(doc, "SAST en CI: obligatorio en todos los repositorios.")
    add_paragraph(doc, "DAST: obligatorio antes de cada release a producción.")
    add_paragraph(doc, "Escaneo de dependencias: Snyk o Dependabot en todos los repos.")

    add_heading(doc, "4. Política de Licencias Open Source", 2)
    add_table(doc, ["Categoría", "Licencias"], [
        ("Permitidas sin restricción", "MIT, Apache-2.0, BSD-2-Clause, BSD-3-Clause, ISC."),
        ("Requieren revisión previa de Legal", "LGPL, MPL-2.0."),
        ("Prohibidas en código entregable", "GPL-2.0, GPL-3.0, AGPL-3.0."),
    ])

    add_heading(doc, "5. Catálogo de ADRs Vigentes (Selección)", 2)
    add_paragraph(doc, "Nexus mantiene 15 Architecture Decision Records activos. Destacan:")
    add_table(doc, ["ID", "Título", "Estado", "Fecha"], [
        ("ADR-007", "Adopción de event-driven con Kafka para integraciones bancarias",
         "Accepted", "2024-08-12"),
        ("ADR-012", "Migración de monolitos legacy a modular monolith antes de microservicios",
         "Accepted", "2025-03-04"),
        ("ADR-014", "Adopción de PostgreSQL como BD relacional por defecto",
         "Accepted", "2025-09-21"),
    ])

    add_heading(doc, "6. Arquitecturas de Referencia por Dominio", 2)
    add_paragraph(
        doc,
        "Existen tres arquitecturas de referencia mantenidas por el equipo de Arquitectura: "
        "core banking (incluye event sourcing y patrones de conciliación), e-commerce retail "
        "(escalabilidad horizontal y CDN), y HIS hospitalario (interoperabilidad HL7 FHIR y "
        "anonimización de datos sensibles).",
    )

    doc.save(f"{OUTPUT_DIR}/D19_Arquitecturas_Referencia_ADRs.docx")
    print("D19 OK")


# ===================================================================
# D20 — Catálogo de Post-Mortems y Lessons Learned
# ===================================================================
def gen_D20():
    doc = Document()
    header_footer(doc, "Post-Mortems y Lessons Learned")
    add_heading(doc, "Catálogo de Post-Mortems y Lessons Learned", 1)
    doc.add_paragraph(
        "Código: TI-CONOC-001 | Versión: 2.4 | Fecha: Marzo 2026 | "
        "Propietario: CTO + PMO | Clasificación: CONFIDENCIAL"
    )

    add_heading(doc, "1. Política Blameless de Post-Mortems", 2)
    add_paragraph(
        doc,
        "Todo post-mortem en Nexus se realiza bajo política blameless: el objetivo es aprender y "
        "mejorar el sistema, no identificar responsables individuales. Los post-mortems tienen "
        "distribución restringida al equipo del proyecto y al CTO. Compartir un post-mortem con un "
        "cliente requiere aprobación previa de la Gerencia Comercial.",
    )

    add_heading(doc, "2. Plantilla Estándar de Post-Mortem", 2)
    secciones = [
        "Resumen ejecutivo y línea de tiempo.",
        "Impacto cuantificado (clientes afectados, servicios caídos, duración).",
        "Detección y respuesta inicial.",
        "Análisis de causa raíz (5 Whys o Ishikawa).",
        "Qué funcionó bien.",
        "Qué no funcionó.",
        "Acciones correctivas con responsable y fecha.",
        "Acciones preventivas estructurales.",
    ]
    for s in secciones:
        doc.add_paragraph(f"• {s}", style="List Bullet")

    add_heading(doc, "3. Métricas Agregadas del Catálogo (2023–2025)", 2)
    add_paragraph(doc, f"Incidentes mayores documentados en el período: 8.")
    add_table(doc, ["Severidad", "MTTR promedio (horas)"], [
        ("P1", "3.2"),
        ("P2", "8.7"),
        ("P3", "26.4"),
        ("P4", "72.0"),
    ])
    add_paragraph(
        doc,
        "El 22% de los incidentes mayores presenta una causa raíz ya observada en incidentes "
        "anteriores. Este indicador es seguido trimestralmente por el CTO como métrica de "
        "efectividad del aprendizaje organizacional.",
    )

    add_heading(doc, "4. Top 5 Causas Raíz Recurrentes", 2)
    causas = [
        "Cambios en producción sin feature flag.",
        "Falla de capacidad por crecimiento no anticipado.",
        "Configuraciones incorrectas en cutover.",
        "Dependencias de terceros sin contingencia.",
        "Gaps en pruebas de carga.",
    ]
    for i, c in enumerate(causas, 1):
        doc.add_paragraph(f"{i}. {c}")

    add_heading(doc, "5. Acciones Estructurales en Curso", 2)
    add_paragraph(
        doc,
        "El comité de arquitectura mantiene un backlog de acciones de mitigación derivadas del "
        "análisis recurrente de causas raíz, incluyendo: adopción obligatoria de feature flags en "
        "todos los servicios críticos, runbooks de capacity planning trimestrales, checklist "
        "estandarizado de cutover y simulacros de pruebas de carga previo a cada release mayor.",
    )

    doc.save(f"{OUTPUT_DIR}/D20_PostMortems_LessonsLearned.docx")
    print("D20 OK")


# ===================================================================
# D21 — Runbook de Delivery / SDLC Interno
# ===================================================================
def gen_D21():
    doc = Document()
    header_footer(doc, "Runbook de Delivery / SDLC", classification="USO INTERNO")
    add_heading(doc, "Runbook de Delivery — SDLC Interno", 1)
    doc.add_paragraph(
        "Código: TI-PROC-002 | Versión: 4.0 | Fecha: Enero 2026 | "
        "Propietario: CTO + PMO | Clasificación: USO INTERNO"
    )

    add_heading(doc, "1. Fases del Ciclo de Delivery", 2)
    add_table(doc, ["Fase", "Duración", "Criterio de salida"], [
        ("Discovery", "2–4 semanas", "Visión aprobada, backlog inicial, arquitectura macro."),
        ("Design", "2–6 semanas", "ADRs principales firmados, prototipos UX, plan de pruebas."),
        ("Build", "Iterativo (sprints)", "Sprints con Definition of Done cumplido."),
        ("Test", "Incluido en Build", "Cobertura mínima alcanzada, pruebas E2E verdes."),
        ("Deploy", "1–2 semanas", "Cutover ejecutado, smoke tests OK, runbook operativo entregado."),
        ("Hypercare", "4 semanas", "Cero incidentes P1 abiertos, KT a soporte completado."),
    ])

    add_heading(doc, "2. Definition of Done a Tres Niveles", 2)
    add_paragraph(doc, "Historia (story):", bold=True)
    add_paragraph(
        doc,
        "Código revisado por par, cobertura de tests unitarios sobre la regla aceptada, "
        "documentación de API actualizada si aplica.",
    )
    add_paragraph(doc, "Sprint:", bold=True)
    add_paragraph(
        doc,
        "Demo realizada, criterios de aceptación validados, deuda técnica registrada en backlog.",
    )
    add_paragraph(doc, "Release:", bold=True)
    add_paragraph(
        doc,
        "Pruebas E2E verdes, plan de rollback validado, runbook operativo entregado a soporte, "
        "manual de usuario actualizado.",
    )

    add_heading(doc, "3. Ceremonias Agile Obligatorias", 2)
    add_table(doc, ["Ceremonia", "Cadencia", "Duración"], [
        ("Daily", "Diaria", "15 minutos"),
        ("Planning", "Quincenal", "2 horas"),
        ("Review", "Quincenal", "1 hora"),
        ("Retrospectiva", "Quincenal", "1 hora"),
        ("Refinement", "Semanal", "1 hora"),
    ])

    add_heading(doc, "4. Knowledge Transfer y Handoff a Soporte", 2)
    add_paragraph(
        doc,
        "Todo proyecto debe ejecutar un Knowledge Transfer al equipo de soporte con duración "
        "mínima obligatoria de 2 semanas. El KT incluye: walkthrough de arquitectura, runbook "
        "operativo, escenarios típicos de soporte y sesiones de shadowing.",
    )

    add_heading(doc, "5. Artefactos Obligatorios por Fase", 2)
    add_table(doc, ["Fase", "Artefactos"], [
        ("Discovery", "Documento de visión, backlog inicial, mapa de stakeholders."),
        ("Design", "ADRs, prototipos UX, plan de pruebas, modelo de datos."),
        ("Build", "Código en repositorio, tests automatizados, demos quincenales."),
        ("Deploy", "Runbook operativo, plan de rollback, evidencia de smoke tests."),
        ("Hypercare", "Reporte semanal de tickets, KT firmado por soporte, manual de usuario."),
    ])

    doc.save(f"{OUTPUT_DIR}/D21_Runbook_Delivery_SDLC.docx")
    print("D21 OK")


# ===================================================================
# D22 — Gestión de Incidentes con Cliente y SLA
# ===================================================================
def gen_D22():
    doc = Document()
    header_footer(doc, "Gestión de Incidentes con Cliente y SLA")
    add_heading(doc, "Gestión de Incidentes con Cliente y Cumplimiento de SLA", 1)
    doc.add_paragraph(
        "Código: COMERCIAL-PROC-002 | Versión: 3.3 | Fecha: Febrero 2026 | "
        "Propietario: Gerencia Comercial + CTO | Clasificación: CONFIDENCIAL"
    )

    add_heading(doc, "1. Alcance", 2)
    add_paragraph(
        doc,
        "Este procedimiento aplica a incidentes reportados por el cliente sobre servicios o "
        "soluciones entregadas por Nexus. Para incidentes de TI internos referirse al D10 "
        "(Procedimiento de Gestión de Incidentes de TI).",
    )

    add_heading(doc, "2. Clasificación de Severidad y SLA", 2)
    add_table(
        doc,
        ["Severidad", "Descripción", "Respuesta máxima", "Resolución máxima"],
        [
            ("P1", "Servicio crítico caído sin workaround.", "15 minutos", "4 horas"),
            ("P2", "Funcionalidad mayor afectada, workaround disponible.", "30 minutos", "8 horas"),
            ("P3", "Funcionalidad menor afectada.", "2 horas", "24 horas"),
            ("P4", "Consulta o solicitud de mejora.", "24 horas", "5 días hábiles"),
        ],
    )

    add_heading(doc, "3. Penalidades Contractuales", 2)
    add_paragraph(
        doc,
        "Las penalidades por incumplimiento de SLA están definidas en cada contrato individual "
        "(ver D18 — Cartera de Clientes y Maestro de Contratos). Toda penalidad debe ser "
        "registrada en el módulo financiero por la Gerencia Comercial dentro de 5 días hábiles "
        "desde la confirmación del incumplimiento.",
    )

    add_heading(doc, "4. Niveles de Escalamiento (RACI)", 2)
    niveles = ["L1 Soporte", "L2 Ingeniería de Producto", "L3 Arquitectura", "Gerencia"]
    for i, n in enumerate(niveles, 1):
        doc.add_paragraph(f"{i}. {n}")

    add_heading(doc, "5. Activación de War Room", 2)
    add_paragraph(
        doc,
        "Se activa un war room ante toda P1 que supere 1 hora sin avance, o ante una P2 que "
        "supere 4 horas sin avance. El war room está compuesto por: líder del proyecto, arquitecto "
        "asignado, representante de operaciones, sponsor Nexus de la cuenta y, cuando "
        "corresponda, sponsor del cliente.",
    )

    add_heading(doc, "6. Plantillas de Comunicación con el Cliente", 2)
    add_paragraph(doc, "Toda comunicación con el cliente sigue una de cuatro plantillas oficiales:")
    plantillas = [
        "Acuse de recibo del incidente (en el primer contacto).",
        "Actualización periódica de avance (con cadencia según severidad).",
        "Reason for Outage (RFO) preliminar y formal.",
        "Comunicación de cierre con resumen y acciones de seguimiento.",
    ]
    for p in plantillas:
        doc.add_paragraph(f"• {p}", style="List Bullet")

    add_heading(doc, "7. Post-Mortem Obligatorio", 2)
    add_paragraph(
        doc,
        "Todo incidente de severidad P1 o P2 requiere un post-mortem formal (ver D20), entregable "
        "en un máximo de 5 días hábiles desde la resolución del incidente.",
    )

    doc.save(f"{OUTPUT_DIR}/D22_Incidentes_Cliente_SLA.docx")
    print("D22 OK")


# ===================================================================
# D23 — Política de Uso de IA, LLMs y Datos Confidenciales
# ===================================================================
def gen_D23():
    doc = Document()
    header_footer(doc, "Política de Uso de IA y LLMs",
                  classification="USO INTERNO — Lectura obligatoria")
    add_heading(doc, "Política de Uso de IA, LLMs y Datos Confidenciales", 1)
    doc.add_paragraph(
        "Código: TI-POL-002 | Versión: 1.4 | Fecha: Abril 2026 | "
        "Propietario: CISO — Diana Reyes Castañeda + DPO — Mónica Salinas Bustamante | "
        "Clasificación: USO INTERNO"
    )

    add_heading(doc, "1. Niveles de Clasificación de Datos", 2)
    add_table(doc, ["Nivel", "Ejemplos típicos"], [
        ("Público", "Material de marketing, web pública, casos publicados."),
        ("Interno", "Comunicaciones internas, organigrama, políticas employee-facing."),
        ("Confidencial", "Código de clientes, propuestas comerciales, pricing, ADRs."),
        ("Restringido", "Datos personales sensibles, financieros del cliente, credenciales."),
    ])

    add_heading(doc, "2. Matriz de Uso Permitido de LLMs por Nivel de Dato", 2)
    add_table(
        doc,
        ["Nivel", "ChatGPT público", "Claude público", "Copilot", "IA corporativa local Nexus"],
        [
            ("Público", "Sí", "Sí", "Sí", "Sí"),
            ("Interno", "No", "No", "Solo con cuenta enterprise y opt-out de entrenamiento", "Sí"),
            ("Confidencial", "No", "No", "No", "Sí"),
            ("Restringido", "No", "No", "No", "Solo con aprobación caso a caso del CISO"),
        ],
    )

    add_heading(doc, "3. Casos de Uso Aprobados de IA Generativa", 2)
    aprobados = [
        "Generación de boilerplate y código no productivo.",
        "Redacción de comunicaciones internas.",
        "Traducción de documentación pública.",
        "Búsqueda semántica sobre repositorios internos vía el chatbot Nexus local.",
    ]
    for a in aprobados:
        doc.add_paragraph(f"• {a}", style="List Bullet")

    add_heading(doc, "4. Casos de Uso Prohibidos", 2)
    prohibidos = [
        "Subir código de clientes a LLMs públicos.",
        "Ingresar datos personales identificables (DNI, salud, financieros) en herramientas no aprobadas.",
        "Compartir contratos, propuestas comerciales o pricing en LLMs externos.",
        "Usar herramientas con IA embebida no inventariadas por el CISO.",
    ]
    for p in prohibidos:
        doc.add_paragraph(f"• {p}", style="List Bullet")

    add_heading(doc, "5. Proceso de Aprobación de Nuevas Herramientas con IA", 2)
    add_paragraph(
        doc,
        "Toda herramienta con IA embebida (por ejemplo Notion AI, Granola, Cursor, etc.) requiere "
        "aprobación previa del CISO mediante ticket TI-IA. La respuesta se entrega en máximo 5 "
        "días hábiles e incluye revisión del DPA, política de retención y opciones de opt-out de "
        "entrenamiento del proveedor.",
    )

    add_heading(doc, "6. Auditoría y Registro", 2)
    add_paragraph(
        doc,
        "El CISO audita trimestralmente el uso de IA en la organización. Anualmente se entrega un "
        "reporte al Directorio. El uso de la IA corporativa local de Nexus es loggeado de forma "
        "centralizada con propósito de auditoría y mejora continua.",
    )

    add_heading(doc, "7. Sanciones por Incumplimiento", 2)
    add_paragraph(
        doc,
        "El incumplimiento de esta política se trata como uso indebido de información conforme al "
        "D06 (Código de Conducta y Ética Empresarial), con sanciones que pueden ir desde "
        "amonestación escrita hasta desvinculación, según severidad y reincidencia.",
    )

    doc.save(f"{OUTPUT_DIR}/D23_Politica_IA_LLMs.docx")
    print("D23 OK")


# ----- main -----

if __name__ == "__main__":
    print(f"Output dir: {OUTPUT_DIR}")
    gen_D16()
    gen_D17()
    gen_D18()
    gen_D19()
    gen_D20()
    gen_D21()
    gen_D22()
    gen_D23()
    print("Batch Tier-3 (D16-D23) completado.")
