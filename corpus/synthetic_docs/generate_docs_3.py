from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = "/sessions/youthful-brave-goldberg/mnt/14. Local-RAG/corpus/synthetic_docs"

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def header_footer(doc, title):
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = f"Nexus Soluciones S.A.C. — {title}"
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer
    footer.paragraphs[0].text = "Documento Interno — Uso Exclusivo de Colaboradores Nexus Soluciones S.A.C."
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# D11 — Política de Protección de Datos y Manejo de Información de Clientes
def gen_D11():
    doc = Document()
    header_footer(doc, "Política de Protección de Datos")
    add_heading(doc, "Política de Protección de Datos y Manejo de Información de Clientes", 1)
    doc.add_paragraph("Código: LEGAL-POL-002 | Versión: 2.2 | Fecha: Enero 2026 | Propietario: DPO — Mónica Salinas Bustamante")

    add_heading(doc, "1. Marco Legal", 2)
    add_paragraph(doc, "Esta política se rige por la Ley N.° 29733 — Ley de Protección de Datos Personales del Perú y su Reglamento (D.S. N.° 003-2013-JUS). Nexus Soluciones S.A.C. está registrada en el Registro Nacional de Protección de Datos Personales de INDECOPI.")

    add_heading(doc, "2. Oficial de Protección de Datos (DPO)", 2)
    add_paragraph(doc, "El Oficial de Protección de Datos de Nexus Soluciones S.A.C. es:\nMónica Salinas Bustamante\ndpo@nexussoluciones.pe | Ext. 106\nGerencia Legal y Compliance — Piso 8, San Isidro")

    add_heading(doc, "3. Principios de Tratamiento de Datos", 2)
    principios = [
        ("Legalidad", "Los datos solo se tratan con base legal válida: consentimiento, contrato o mandato legal."),
        ("Finalidad", "Los datos se recopilan para fines específicos, explícitos y legítimos."),
        ("Proporcionalidad", "Solo se recaban los datos estrictamente necesarios para el fin declarado."),
        ("Calidad", "Los datos deben ser exactos, actualizados y completos."),
        ("Seguridad", "Se aplican medidas técnicas y organizativas para proteger los datos."),
        ("Confidencialidad", "Los datos no se comunican a terceros sin autorización o base legal."),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Principio"
    t.rows[0].cells[1].text = "Descripción"
    for nombre, desc in principios:
        row = t.add_row()
        row.cells[0].text = nombre
        row.cells[1].text = desc

    add_heading(doc, "4. Manejo de Datos de Clientes en Proyectos", 2)
    reglas = [
        "Los datos de clientes solo pueden almacenarse en la infraestructura autorizada por el propio cliente (on-premise del cliente o cloud acordado contractualmente).",
        "Está prohibido copiar datos de clientes a dispositivos personales, unidades USB, correo personal o servicios en la nube no autorizados.",
        "El acceso a datos de clientes debe ser aprobado por el jefe de proyecto y registrado en Jira.",
        "Al cierre del contrato, todos los datos del cliente deben eliminarse dentro de los 30 días siguientes a la fecha de cierre, documentando el proceso con un acta de destrucción firmada por el DPO.",
    ]
    for r in reglas:
        doc.add_paragraph(f"• {r}", style="List Bullet")

    add_heading(doc, "5. Protocolo ante una Brecha de Datos", 2)
    add_paragraph(doc, "En caso de sospecha o confirmación de brecha de datos que involucre información personal:")
    proto = [
        "El colaborador reporta inmediatamente al CISO (d.reyes@nexussoluciones.pe) y al DPO (dpo@nexussoluciones.pe).",
        "El DPO evalúa el riesgo para los titulares de datos en un plazo máximo de 24 horas.",
        "Si existe riesgo alto, el DPO notifica a INDECOPI y al cliente afectado en un plazo máximo de 72 horas desde la detección.",
        "TI documenta el incidente con un informe post-brecha en 5 días hábiles.",
        "Si aplica, se inician acciones legales y correcciones al sistema de gestión ISO 27001.",
    ]
    for i, p in enumerate(proto, 1):
        doc.add_paragraph(f"{i}. {p}")

    add_heading(doc, "6. Derechos de los Titulares de Datos", 2)
    derechos = ["Acceso a sus datos", "Rectificación de datos incorrectos", "Cancelación o supresión", "Oposición al tratamiento"]
    for d in derechos:
        doc.add_paragraph(f"• {d}", style="List Bullet")
    add_paragraph(doc, "Las solicitudes de ejercicio de derechos deben dirigirse a dpo@nexussoluciones.pe. El plazo de respuesta es de 20 días hábiles.")

    doc.save(f"{OUTPUT_DIR}/D11_Proteccion_Datos_Clientes.docx")
    print("D11 OK")

# D12 — Procedimiento de Compras y Gestión de Proveedores
def gen_D12():
    doc = Document()
    header_footer(doc, "Procedimiento de Compras y Gestión de Proveedores")
    add_heading(doc, "Procedimiento de Compras y Gestión de Proveedores", 1)
    doc.add_paragraph("Código: ADMIN-PROC-001 | Versión: 3.0 | Fecha: Enero 2026 | Propietario: Gerencia de Finanzas y Administración")

    add_heading(doc, "1. Objetivo", 2)
    add_paragraph(doc, "Establecer los lineamientos para la adquisición de bienes y servicios por parte de Nexus Soluciones S.A.C., asegurando eficiencia, transparencia y control del gasto.")

    add_heading(doc, "2. Matriz de Aprobación por Monto", 2)
    montos = [
        ("Hasta S/. 500", "Jefe de área", "No requerida", "Solicitud verbal o email + ticket en Jira"),
        ("S/. 501 – S/. 2,000", "Jefe de área + Analista de Compras", "Sí", "Formulario intranet + 2 cotizaciones"),
        ("S/. 2,001 – S/. 10,000", "CFO — Patricia Quispe", "Sí", "Formulario intranet + 3 cotizaciones + sustento técnico"),
        ("Mayor a S/. 10,000", "Gerente General — Ricardo Mendoza", "Sí", "Proceso de licitación + informe de evaluación + comité"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["Rango", "Aprobador", "Orden de Compra", "Proceso"]):
        t.rows[0].cells[i].text = h
    for row_data in montos:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    add_heading(doc, "3. Proceso de Compra", 2)
    pasos = [
        "El área solicitante identifica la necesidad y accede al formulario en intranet.nexussoluciones.pe/compras.",
        "Se solicitan cotizaciones al número requerido según el rango de monto.",
        "El Analista de Compras (Carmen Villacorta Ruiz — c.villacorta@nexussoluciones.pe, Ext. 602) evalúa las cotizaciones.",
        "La solicitud se envía al aprobador correspondiente según la Matriz de Aprobación.",
        "El aprobador autoriza o rechaza en Jira con comentario justificado.",
        "Compras emite la Orden de Compra (si aplica) y coordina con el proveedor.",
        "El área solicitante confirma la recepción del bien o servicio en el formulario de conformidad.",
        "Finanzas procesa el pago según los términos acordados.",
    ]
    for i, p in enumerate(pasos, 1):
        doc.add_paragraph(f"{i}. {p}")

    add_heading(doc, "4. Proveedores Homologados", 2)
    add_paragraph(doc, "Nexus mantiene un registro actualizado de proveedores homologados en intranet.nexussoluciones.pe/proveedores-homologados. Para contratar con un proveedor no homologado se requiere aprobación del CFO y una evaluación de proveedor previa. La homologación incluye verificación de RUC activo, historial tributario y referencias comerciales.")

    add_heading(doc, "5. Condiciones de Pago", 2)
    add_paragraph(doc, "El plazo estándar de pago a proveedores es de 30 días desde la fecha de factura. Los proveedores que acepten pago a 15 días reciben un descuento por pronto pago del 2%. Los pagos son procesados los martes y jueves de cada semana por el área de Finanzas.")

    add_heading(doc, "6. Contacto del Área de Compras", 2)
    add_paragraph(doc, "Analista de Compras: Carmen Villacorta Ruiz\nc.villacorta@nexussoluciones.pe | Ext. 602\nCaja y Viáticos: Daniela Soto Peña\nd.soto@nexussoluciones.pe | Ext. 603\nJefe de Contabilidad: Roberto Farfán Lezama\nr.farfan@nexussoluciones.pe | Ext. 601")

    doc.save(f"{OUTPUT_DIR}/D12_Compras_Proveedores.docx")
    print("D12 OK")

# D13 — Política de Gastos, Viáticos y Reembolsos
def gen_D13():
    doc = Document()
    header_footer(doc, "Política de Gastos, Viáticos y Reembolsos")
    add_heading(doc, "Política de Gastos, Viáticos y Reembolsos", 1)
    doc.add_paragraph("Código: ADMIN-POL-002 | Versión: 2.5 | Fecha: Enero 2026 | Propietario: Gerencia de Finanzas y Administración")

    add_heading(doc, "1. Tarifas de Viáticos Diarios", 2)
    tarifas = [
        ("Lima (comida y movilidad)", "S/. 35 por día"),
        ("Provincias (alojamiento + alimentación + movilidad)", "S/. 120 por día"),
        ("Internacional (alojamiento + alimentación + movilidad)", "USD 80 por día"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Destino"
    t.rows[0].cells[1].text = "Tarifa"
    for dest, tarifa in tarifas:
        row = t.add_row()
        row.cells[0].text = dest
        row.cells[1].text = tarifa

    add_heading(doc, "2. Transporte", 2)
    add_paragraph(doc, "Taxis y aplicaciones:", bold=True)
    add_paragraph(doc, "Se autoriza el uso de Uber y Cabify con la cuenta corporativa de Nexus. El gasto máximo por viaje es de S/. 60. No se reembolsan taxis no documentados.")
    add_paragraph(doc, "Vuelos:", bold=True)
    add_paragraph(doc, "Siempre en clase económica. Solo se autoriza clase ejecutiva en vuelos de duración superior a 8 horas, previa aprobación del Gerente General.")
    add_paragraph(doc, "Alojamiento:", bold=True)
    aloj = [
        ("Lima", "S/. 280 por noche"),
        ("Provincias", "S/. 220 por noche"),
        ("Internacional", "USD 150 por noche"),
    ]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Destino"
    t2.rows[0].cells[1].text = "Máximo por noche"
    for dest, max_val in aloj:
        row = t2.add_row()
        row.cells[0].text = dest
        row.cells[1].text = max_val

    add_heading(doc, "3. Proceso de Reembolso", 2)
    pasos = [
        "El colaborador ingresa el formulario de reembolso en intranet.nexussoluciones.pe/reembolsos dentro de los 5 días hábiles siguientes al viaje o gasto.",
        "Se adjuntan todos los comprobantes de pago (boletas, facturas o recibos electrónicos).",
        "El jefe directo aprueba el formulario en Jira.",
        "Finanzas valida los montos contra las tarifas de esta política.",
        "El reembolso se abona en la cuenta bancaria registrada del colaborador en un plazo de 7 días hábiles desde la aprobación final.",
    ]
    for i, p in enumerate(pasos, 1):
        doc.add_paragraph(f"{i}. {p}")

    add_heading(doc, "4. Anticipo de Viaje", 2)
    add_paragraph(doc, "El colaborador puede solicitar un anticipo de hasta S/. 500 con 3 días hábiles de anticipación al viaje, a través del formulario en intranet.nexussoluciones.pe/reembolsos sección 'Anticipo'. La liquidación del anticipo debe realizarse dentro de los 3 días hábiles siguientes al retorno. La no liquidación oportuna bloquea futuros anticipos y puede resultar en descuento de haberes.")

    add_heading(doc, "5. Gastos No Reembolsables", 2)
    no_reemb = ["Bebidas alcohólicas", "Entretenimiento personal (cine, spa, etc.)", "Gastos de acompañantes no autorizados", "Multas de tránsito o estacionamiento", "Gastos sin comprobante de pago válido", "Gastos que superen los topes establecidos en esta política sin aprobación previa"]
    for g in no_reemb:
        doc.add_paragraph(f"• {g}", style="List Bullet")

    add_heading(doc, "6. Contacto", 2)
    add_paragraph(doc, "Caja y Viáticos: Daniela Soto Peña — d.soto@nexussoluciones.pe | Ext. 603")

    doc.save(f"{OUTPUT_DIR}/D13_Gastos_Viaticos_Reembolsos.docx")
    print("D13 OK")

# D14 — Reglamento de Uso de Instalaciones y Espacios Corporativos
def gen_D14():
    doc = Document()
    header_footer(doc, "Reglamento de Uso de Instalaciones y Espacios Corporativos")
    add_heading(doc, "Reglamento de Uso de Instalaciones y Espacios Corporativos", 1)
    doc.add_paragraph("Código: ADMIN-REG-001 | Versión: 2.0 | Fecha: Enero 2026 | Propietario: Gerencia de Finanzas y Administración")

    add_heading(doc, "1. Salas de Reuniones — Sede San Isidro", 2)
    salas = [
        ("Sala Arequipa", "8", "8", "TV 65', videoconferencia, pizarra"),
        ("Sala Cusco", "8", "12", "Proyector, videoconferencia, pizarra"),
        ("Sala Lima", "9", "20", "Proyector 4K, audio, videoconferencia, pizarra doble"),
        ("Sala Trujillo", "9", "6", "TV 55', pizarra"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["Sala", "Piso", "Capacidad", "Equipamiento"]):
        t.rows[0].cells[i].text = h
    for row_data in salas:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    add_heading(doc, "2. Reserva de Salas", 2)
    add_paragraph(doc, "Las salas se reservan a través del calendario de recursos en Microsoft Outlook o Teams. La reserva máxima anticipada es de 14 días. La duración máxima por reserva es de 4 horas continuas. Si la sala no es utilizada en los primeros 15 minutos de la hora de inicio, cualquier colaborador puede ocuparla.")

    add_heading(doc, "3. Cafetería", 2)
    add_paragraph(doc, "Disponible en la sede San Isidro. Horario: 07:30 – 19:00, lunes a viernes. La empresa subvenciona el 50% del consumo. El pago se realiza con la credencial corporativa (descuento de planilla quincenal). La cafetería no está disponible los fines de semana.")

    add_heading(doc, "4. Gimnasio", 2)
    add_paragraph(doc, "El gimnasio corporativo está disponible en San Isidro con los siguientes horarios:\n• Mañana: 06:30 – 08:30, lunes a viernes\n• Tarde: 18:00 – 20:00, lunes a viernes\nLa empresa otorga una subvención de S/. 80 mensuales para uso del gimnasio (ya sea el corporativo u otros con convenio). Consultar la lista de gimnasios con convenio en intranet.nexussoluciones.pe/beneficios.")

    add_heading(doc, "5. Estacionamiento", 2)
    add_paragraph(doc, "El estacionamiento en San Isidro se asigna por antigüedad y rol. La lista es administrada por el área de Administración. Los colaboradores sin plaza asignada pueden usar estacionamiento en las inmediaciones bajo su costo. Las motocicletas tienen zona designada en B1.")

    add_heading(doc, "6. Lockers", 2)
    add_paragraph(doc, "Los lockers están disponibles en los pisos 8 y 9. Son asignados por RR.HH. al ingreso. El colaborador es responsable de su uso adecuado. Al cese, debe entregar el locker vacío y la llave a RR.HH.")

    add_heading(doc, "7. Política de Visitas", 2)
    add_paragraph(doc, "Todos los visitantes externos deben:")
    visitas = [
        "Registrarse en recepción presentando su DNI o pasaporte.",
        "Portar la credencial de visitante durante toda su permanencia.",
        "Ser acompañados en todo momento por el colaborador anfitrión.",
        "No acceder a zonas restringidas: Sala de Servidores (Piso 9) ni Archivo Confidencial (Piso 8).",
    ]
    for v in visitas:
        doc.add_paragraph(f"• {v}", style="List Bullet")

    add_heading(doc, "8. Zonas Restringidas", 2)
    add_paragraph(doc, "Las siguientes zonas requieren acceso especial autorizado por TI o Administración:\n• Sala de Servidores — Piso 9, San Isidro: acceso solo para personal de TI autorizado.\n• Archivo Confidencial — Piso 8, San Isidro: acceso solo para Legal y alta gerencia.")

    doc.save(f"{OUTPUT_DIR}/D14_Reglamento_Instalaciones.docx")
    print("D14 OK")

# D15 — Directorio Corporativo y Organigrama
def gen_D15():
    doc = Document()
    header_footer(doc, "Directorio Corporativo y Organigrama")
    add_heading(doc, "Directorio Corporativo y Organigrama", 1)
    doc.add_paragraph("Versión: Enero 2026 | Responsable: Gerencia de RR.HH. | Actualización: Trimestral")

    add_heading(doc, "1. Datos Generales de la Empresa", 2)
    datos = [
        ("Razón Social", "Nexus Soluciones S.A.C."),
        ("RUC", "20601847392"),
        ("Fundación", "14 de marzo de 2017"),
        ("Sector", "Tecnología de la Información y Consultoría"),
        ("Teléfono Central", "(01) 611-4200"),
        ("Dominio de Correo", "@nexussoluciones.pe"),
        ("Web", "www.nexussoluciones.pe"),
        ("Sede Principal", "Av. Javier Prado Este 4200, Pisos 8-9, San Isidro, Lima"),
        ("Oficina de Proyectos", "Calle Berlín 1044, Piso 3, Miraflores, Lima"),
        ("Total Colaboradores", "148 (132 en planilla, 16 locadores)"),
        ("Certificaciones", "ISO 9001:2015, ISO 27001:2022"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for campo, valor in datos:
        row = t.add_row()
        row.cells[0].text = campo
        row.cells[1].text = valor

    add_heading(doc, "2. Alta Dirección", 2)
    alta = [
        ("Presidente del Directorio", "Eduardo Vargas Llona", "e.vargas@nexussoluciones.pe", "—"),
        ("Gerente General (CEO)", "Ricardo Mendoza Paredes", "r.mendoza@nexussoluciones.pe", "101"),
        ("Gerente de Finanzas y Administración (CFO)", "Patricia Quispe Huamán", "p.quispe@nexussoluciones.pe", "102"),
        ("Gerente de Tecnología (CTO)", "Andrés Flores Castillo", "a.flores@nexussoluciones.pe", "103"),
        ("Gerente de Recursos Humanos (CHRO)", "Silvia Torres Vega", "s.torres@nexussoluciones.pe", "104"),
        ("Gerente Comercial y de Proyectos", "Carlos Amat y León Ríos", "c.amat@nexussoluciones.pe", "105"),
        ("Gerente Legal y Compliance", "Mónica Salinas Bustamante", "m.salinas@nexussoluciones.pe", "106"),
    ]
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    for i, h in enumerate(["Cargo", "Nombre", "Correo", "Ext."]):
        t2.rows[0].cells[i].text = h
    for row_data in alta:
        row = t2.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    add_heading(doc, "3. Directorio por Departamento", 2)
    departamentos = [
        ("Recursos Humanos (8 personas)", [
            ("Gerente RRHH", "Silvia Torres Vega", "s.torres@nexussoluciones.pe", "104"),
            ("Especialista de Selección", "Lucía Romero Díaz", "l.romero@nexussoluciones.pe", "201"),
            ("Especialista de Compensaciones", "Jorge Palacios Ríos", "j.palacios@nexussoluciones.pe", "202"),
            ("Analista de Bienestar", "Valeria Chávez Núñez", "v.chavez@nexussoluciones.pe", "203"),
        ]),
        ("Tecnología e Infraestructura (28 personas)", [
            ("CTO / Gerente TI", "Andrés Flores Castillo", "a.flores@nexussoluciones.pe", "103"),
            ("Líder de Infraestructura", "Marco Huanca Soto", "m.huanca@nexussoluciones.pe", "301"),
            ("CISO", "Diana Reyes Castañeda", "d.reyes@nexussoluciones.pe", "302"),
            ("Mesa de Ayuda (Helpdesk)", "—", "helpdesk@nexussoluciones.pe", "300"),
        ]),
        ("Desarrollo de Software (42 personas)", [
            ("Tech Lead Backend", "Sebastián Quiroga Medina", "s.quiroga@nexussoluciones.pe", "401"),
            ("Tech Lead Frontend", "Andrea Villanueva Cruz", "a.villanueva@nexussoluciones.pe", "402"),
            ("QA Lead", "Fernando Cáceres Alva", "f.caceres@nexussoluciones.pe", "403"),
        ]),
        ("Consultoría y Proyectos (35 personas)", [
            ("Gerente Comercial", "Carlos Amat y León Ríos", "c.amat@nexussoluciones.pe", "105"),
            ("PMO Lead", "Rosa Mendívil Parodi", "r.mendivil@nexussoluciones.pe", "501"),
            ("Consultor Senior", "Óscar Tapia Gutiérrez", "o.tapia@nexussoluciones.pe", "502"),
        ]),
        ("Finanzas y Administración (14 personas)", [
            ("CFO", "Patricia Quispe Huamán", "p.quispe@nexussoluciones.pe", "102"),
            ("Jefe de Contabilidad", "Roberto Farfán Lezama", "r.farfan@nexussoluciones.pe", "601"),
            ("Analista de Compras", "Carmen Villacorta Ruiz", "c.villacorta@nexussoluciones.pe", "602"),
            ("Caja y Viáticos", "Daniela Soto Peña", "d.soto@nexussoluciones.pe", "603"),
        ]),
        ("Legal y Compliance (5 personas)", [
            ("Gerente Legal / DPO", "Mónica Salinas Bustamante", "m.salinas@nexussoluciones.pe", "106"),
            ("Abogado Senior", "Hernán Delgado Espino", "h.delgado@nexussoluciones.pe", "701"),
            ("Contacto DPO", "—", "dpo@nexussoluciones.pe", "106"),
        ]),
        ("Ventas y Marketing (11 personas)", [
            ("Gerente Comercial", "Carlos Amat y León Ríos", "c.amat@nexussoluciones.pe", "105"),
            ("Key Account Manager", "Fiorella Montoya Bravo", "f.montoya@nexussoluciones.pe", "801"),
            ("Marketing Digital", "Diego Cornejo Suárez", "d.cornejo@nexussoluciones.pe", "802"),
        ]),
    ]
    for dept_nombre, contactos in departamentos:
        add_paragraph(doc, dept_nombre, bold=True)
        t_dept = doc.add_table(rows=1, cols=4)
        t_dept.style = "Table Grid"
        for i, h in enumerate(["Cargo", "Nombre", "Correo", "Ext."]):
            t_dept.rows[0].cells[i].text = h
        for row_data in contactos:
            row = t_dept.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = val
        doc.add_paragraph("")

    add_heading(doc, "4. Contactos de Emergencia y Canales Corporativos", 2)
    emergencia = [
        ("Canal de Denuncias Ética", "denuncias@nexussoluciones.pe | Línea: 0800-00-555"),
        ("Oficial de Protección de Datos", "dpo@nexussoluciones.pe"),
        ("Mesa de Ayuda TI", "helpdesk@nexussoluciones.pe | Ext. 300"),
        ("Recepción San Isidro", "(01) 611-4200 | Ext. 100"),
        ("Emergencias internas SST", "m.huanca@nexussoluciones.pe | Ext. 301"),
    ]
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Canal"
    t3.rows[0].cells[1].text = "Contacto"
    for canal, contacto in emergencia:
        row = t3.add_row()
        row.cells[0].text = canal
        row.cells[1].text = contacto

    doc.save(f"{OUTPUT_DIR}/D15_Directorio_Organigrama.docx")
    print("D15 OK")

gen_D11()
gen_D12()
gen_D13()
gen_D14()
gen_D15()
print("Batch 3 (D11-D15) completado.")
