from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/sessions/youthful-brave-goldberg/mnt/14. Local-RAG/corpus/synthetic_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
    return row

def header_footer(doc, title):
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = f"Nexus Soluciones S.A.C. — {title}"
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer
    footer.paragraphs[0].text = "Documento Interno — Uso Exclusivo de Colaboradores Nexus Soluciones S.A.C."
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─────────────────────────────────────────────
# D01 — Manual de Bienvenida
# ─────────────────────────────────────────────
def gen_D01():
    doc = Document()
    header_footer(doc, "Manual de Bienvenida")
    add_heading(doc, "Manual de Bienvenida", 1)
    add_heading(doc, "Nexus Soluciones S.A.C.", 2)
    doc.add_paragraph("Versión 3.1 | Fecha de emisión: Enero 2026 | Responsable: Gerencia de Recursos Humanos")
    doc.add_paragraph("")

    add_heading(doc, "1. Mensaje de Bienvenida del Gerente General", 2)
    add_paragraph(doc, "Estimado/a colaborador/a,")
    add_paragraph(doc, "En nombre de todo el equipo de Nexus Soluciones S.A.C., te damos la más cordial bienvenida. Desde nuestra fundación el 14 de marzo de 2017, hemos construido una empresa donde la excelencia técnica y la integridad van de la mano. Hoy formas parte de un equipo de 148 personas comprometidas con transformar digitalmente a nuestros clientes, garantizando siempre la confidencialidad e integridad de su información.")
    add_paragraph(doc, "Confiamos en tu talento y esperamos que este sea el inicio de una relación larga y exitosa.")
    add_paragraph(doc, "Ricardo Mendoza Paredes\nGerente General\nr.mendoza@nexussoluciones.pe | Ext. 101")

    add_heading(doc, "2. Quiénes Somos", 2)
    add_paragraph(doc, "Nexus Soluciones S.A.C. (RUC 20601847392) es una empresa peruana especializada en desarrollo de software a medida, consultoría tecnológica e infraestructura cloud para el sector financiero, retail y salud. Contamos con certificaciones ISO 9001:2015 e ISO 27001:2022, lo que refleja nuestro compromiso con la calidad y la seguridad de la información.")

    add_heading(doc, "3. Misión, Visión y Valores", 2)
    add_paragraph(doc, "Misión:", bold=True)
    add_paragraph(doc, "Transformar digitalmente a nuestros clientes mediante soluciones tecnológicas de alta calidad, asegurando la confidencialidad e integridad de su información.")
    add_paragraph(doc, "Visión:", bold=True)
    add_paragraph(doc, "Ser la empresa de tecnología de referencia en el Perú para sectores con alta regulación de datos al año 2028.")
    add_paragraph(doc, "Valores:", bold=True)
    for v in ["Integridad", "Innovación", "Compromiso", "Confidencialidad", "Excelencia"]:
        doc.add_paragraph(f"• {v}", style="List Bullet")

    add_heading(doc, "4. Nuestras Sedes", 2)
    add_paragraph(doc, "Sede Principal — San Isidro:")
    add_paragraph(doc, "Av. Javier Prado Este 4200, Pisos 8 y 9, San Isidro, Lima. Teléfono central: (01) 611-4200. Capacidad: 130 personas. Cuenta con estacionamiento.")
    add_paragraph(doc, "Oficina de Proyectos — Miraflores:")
    add_paragraph(doc, "Calle Berlín 1044, Piso 3, Miraflores, Lima. Capacidad: 40 personas.")

    add_heading(doc, "5. Estructura Organizacional", 2)
    add_paragraph(doc, "Nexus cuenta con 7 áreas principales: Recursos Humanos (8 personas), Tecnología e Infraestructura (28), Desarrollo de Software (42), Consultoría y Proyectos (35), Finanzas y Administración (14), Legal y Compliance (5), y Ventas y Marketing (11).")

    add_heading(doc, "6. Beneficios para Colaboradores", 2)
    beneficios = [
        "30 días de vacaciones al año",
        "EPS RIMAC con cobertura 70% empresa / 30% colaborador",
        "Seguro de vida",
        "Subvención de gimnasio: S/. 80 mensuales",
        "Día libre por cumpleaños",
        "2 sesiones mensuales de apoyo psicológico gratuitas",
        "Cafetería subvencionada al 50% en San Isidro",
        "Modalidad híbrida: 2 días de trabajo remoto por semana (desde el 3.er mes)",
        "Clases de inglés gratuitas: martes y jueves 7:00–8:00 am (virtual)",
        "Bono de desempeño anual: hasta 1.5 sueldos"
    ]
    for b in beneficios:
        doc.add_paragraph(f"• {b}", style="List Bullet")

    add_heading(doc, "7. Horarios de Trabajo", 2)
    add_paragraph(doc, "La jornada estándar es de 48 horas semanales, 8 horas diarias. Contamos con tres turnos:")
    turnos = [
        ("Turno A", "08:00 – 17:30", "1 hora de almuerzo"),
        ("Turno B", "09:00 – 18:30", "1 hora de almuerzo"),
        ("Turno C (solo Miraflores)", "10:00 – 19:30", "1 hora de almuerzo"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Turno"
    t.rows[0].cells[1].text = "Horario"
    t.rows[0].cells[2].text = "Almuerzo"
    for nombre, horario, almuerzo in turnos:
        row = t.add_row()
        row.cells[0].text = nombre
        row.cells[1].text = horario
        row.cells[2].text = almuerzo

    add_heading(doc, "8. Contactos Clave para tu Primer Día", 2)
    contactos = [
        ("Recursos Humanos", "Silvia Torres Vega", "s.torres@nexussoluciones.pe", "104"),
        ("Mesa de Ayuda TI", "Helpdesk", "helpdesk@nexussoluciones.pe", "300"),
        ("Recepción San Isidro", "—", "(01) 611-4200", "100"),
    ]
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    for i, h in enumerate(["Área", "Contacto", "Email", "Ext."]):
        t2.rows[0].cells[i].text = h
    for area, nombre, email, ext in contactos:
        row = t2.add_row()
        row.cells[0].text = area
        row.cells[1].text = nombre
        row.cells[2].text = email
        row.cells[3].text = ext

    doc.save(f"{OUTPUT_DIR}/D01_Manual_Bienvenida.docx")
    print("D01 OK")

# ─────────────────────────────────────────────
# D02 — Procedimiento de Onboarding
# ─────────────────────────────────────────────
def gen_D02():
    doc = Document()
    header_footer(doc, "Procedimiento de Onboarding")
    add_heading(doc, "Procedimiento de Onboarding de Nuevos Colaboradores", 1)
    doc.add_paragraph("Código: RRHH-PROC-001 | Versión: 2.4 | Fecha: Marzo 2026 | Propietario: Gerencia de RR.HH.")

    add_heading(doc, "1. Objetivo", 2)
    add_paragraph(doc, "Establecer el proceso estándar para la incorporación efectiva de nuevos colaboradores a Nexus Soluciones S.A.C., garantizando que cuenten con los accesos, equipos, información y contexto necesarios para iniciar sus funciones de manera productiva.")

    add_heading(doc, "2. Alcance", 2)
    add_paragraph(doc, "Aplica a todos los colaboradores que ingresen a planilla de Nexus Soluciones S.A.C. en cualquiera de sus sedes. El proceso de onboarding tiene una duración de 15 días calendario.")

    add_heading(doc, "3. Responsables", 2)
    resp = [
        ("Gerencia de RR.HH.", "Silvia Torres Vega — s.torres@nexussoluciones.pe", "Coordinación general del proceso, firma de contrato y NDA, asignación de buddy"),
        ("Gerencia de TI", "Andrés Flores Castillo — a.flores@nexussoluciones.pe", "Provisión de equipos y credenciales de acceso"),
        ("Jefe Directo", "Según área asignada", "Inducción técnica, presentación de equipo, asignación de primeras tareas"),
        ("Buddy Asignado", "Colega del área", "Acompañamiento durante los primeros 30 días"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["Rol", "Responsable", "Actividad Principal"]):
        t.rows[0].cells[i].text = h
    for rol, resp_nombre, actividad in resp:
        row = t.add_row()
        row.cells[0].text = rol
        row.cells[1].text = resp_nombre
        row.cells[2].text = actividad

    add_heading(doc, "4. Cronograma del Proceso", 2)
    etapas = [
        ("Día 1", [
            "Recepción por RR.HH. — Silvia Torres Vega o delegado (8:00 am, recepción San Isidro)",
            "Firma de contrato de trabajo y Acuerdo de No Divulgación (NDA)",
            "Entrega de equipos: laptop corporativa (Lenovo ThinkPad E14 Gen 4 o Dell Latitude 5540), teléfono IP, credencial de acceso físico y kit de bienvenida Nexus",
            "Creación de cuenta corporativa: formato nombre.apellido@nexussoluciones.pe",
            "Activación de MFA (Microsoft Authenticator) — obligatorio",
            "Recorrido guiado por instalaciones: ubicación de salas, cafetería, salidas de emergencia y puntos de reunión",
        ]),
        ("Días 2–5", [
            "Inducción corporativa: misión, visión, valores, certificaciones ISO, organigrama y cultura Nexus",
            "Inducción técnica con el área: herramientas, metodologías y proyectos vigentes",
            "Capacitación en herramientas internas: Jira, Confluence, GitLab, Microsoft Teams",
            "Presentación formal con el equipo directo y con la Gerencia de área",
            "Revisión del Código de Conducta y políticas de seguridad de la información",
        ]),
        ("Días 6–10", [
            "Shadowing con el buddy asignado: observación de flujos de trabajo reales",
            "Activación de acceso completo a sistemas según perfil de usuario aprobado por jefe directo",
            "Reunión 1:1 con Gerente de área (30 minutos mínimo)",
            "Acceso a Nexus Academy: plataforma de e-learning en intranet.nexussoluciones.pe/academy",
        ]),
        ("Días 11–15", [
            "Asignación de primeras tareas bajo supervisión del jefe directo",
            "Evaluación de inducción: cuestionario en intranet.nexussoluciones.pe/onboarding-eval",
            "Reunión de cierre con RR.HH.: feedback del proceso, dudas y próximos pasos",
            "Confirmación de período de prueba: 3 meses desde fecha de ingreso",
        ]),
    ]
    for etapa, actividades in etapas:
        add_paragraph(doc, etapa, bold=True)
        for a in actividades:
            doc.add_paragraph(f"• {a}", style="List Bullet")

    add_heading(doc, "5. Programa de Buddy", 2)
    add_paragraph(doc, "Cada colaborador nuevo tiene asignado un buddy del mismo departamento durante los primeros 30 días. El buddy es designado por el jefe directo y debe tener más de 6 meses en la empresa. Sus responsabilidades incluyen responder preguntas del día a día, presentar al equipo extendido y brindar contexto sobre la cultura del área.")

    add_heading(doc, "6. Equipos Entregados en el Primer Día", 2)
    equipos = ["Laptop corporativa (Lenovo ThinkPad E14 Gen 4 o Dell Latitude 5540, imagen Windows 11 Pro)", "Teléfono IP de escritorio", "Credencial de acceso físico (tarjeta magnética)", "Kit de bienvenida Nexus (cuaderno, lapicero, lanyard corporativo)"]
    for e in equipos:
        doc.add_paragraph(f"• {e}", style="List Bullet")
    add_paragraph(doc, "Nota: El colaborador firma un acta de recepción de equipos en el Día 1. Es responsable de los equipos asignados durante toda su permanencia en la empresa.")

    add_heading(doc, "7. Período de Prueba", 2)
    add_paragraph(doc, "El período de prueba es de 3 meses a partir de la fecha de ingreso. Durante este período, tanto el colaborador como la empresa pueden dar por concluida la relación laboral sin expresión de causa. Al finalizar el período de prueba, RR.HH. enviará al jefe directo un formulario de confirmación de colaborador.")

    doc.save(f"{OUTPUT_DIR}/D02_Procedimiento_Onboarding.docx")
    print("D02 OK")

# ─────────────────────────────────────────────
# D03 — Política de Vacaciones, Licencias y Permisos
# ─────────────────────────────────────────────
def gen_D03():
    doc = Document()
    header_footer(doc, "Política de Vacaciones, Licencias y Permisos")
    add_heading(doc, "Política de Vacaciones, Licencias y Permisos", 1)
    doc.add_paragraph("Código: RRHH-POL-002 | Versión: 4.0 | Fecha: Enero 2026 | Propietario: Gerencia de RR.HH.")

    add_heading(doc, "1. Vacaciones Anuales", 2)
    add_paragraph(doc, "Todo colaborador de Nexus Soluciones S.A.C. tiene derecho a 30 días calendario de vacaciones remuneradas por cada año completo de servicios, conforme a la legislación laboral peruana vigente.")

    add_heading(doc, "1.1 Condiciones de Goce", 3)
    cond = [
        "El período mínimo continuo de vacaciones es de 7 días calendario.",
        "Las vacaciones deben solicitarse con un mínimo de 15 días de anticipación.",
        "La aprobación requiere conformidad del jefe directo y registro en el Portal de RR.HH.: intranet.nexussoluciones.pe/vacaciones.",
        "Las vacaciones se abonan con la remuneración mensual completa correspondiente a los 30 días.",
        "Se pueden acumular hasta 15 días no gozados, que deben utilizarse antes del 31 de marzo del año siguiente. Los días no utilizados más allá de ese límite se abonan como vacaciones truncas.",
    ]
    for c in cond:
        doc.add_paragraph(f"• {c}", style="List Bullet")

    add_heading(doc, "1.2 Proceso de Solicitud", 3)
    pasos = [
        "Ingresar a intranet.nexussoluciones.pe/vacaciones con tus credenciales corporativas.",
        "Seleccionar las fechas deseadas y verificar el saldo disponible.",
        "Enviar solicitud: el sistema notifica automáticamente al jefe directo.",
        "El jefe directo aprueba o rechaza en un plazo máximo de 5 días hábiles.",
        "RR.HH. confirma el registro y notifica a planilla para el procesamiento del pago.",
    ]
    for i, p in enumerate(pasos, 1):
        doc.add_paragraph(f"{i}. {p}")

    add_heading(doc, "2. Licencias", 2)
    licencias = [
        ("Maternidad", "98 días", "Sí", "49 días pre parto y 49 días post parto. Requiere certificado médico."),
        ("Paternidad", "10 días", "Sí", "Desde el día del nacimiento. Presentar partida de nacimiento."),
        ("Matrimonio", "5 días", "Sí", "Dentro de los 15 días calendario de celebrado el matrimonio."),
        ("Fallecimiento familiar directo", "5 días", "Sí", "Cónyuge, hijos o padres. Presentar partida de defunción."),
        ("Fallecimiento familiar indirecto", "3 días", "Sí", "Hermanos o abuelos. Presentar partida de defunción."),
        ("Enfermedad con certificado médico", "Según indicación", "Sí", "Certificado de ESSALUD o médico autorizado. Desde el 4.° día cubre ESSALUD."),
        ("Adopción", "30 días", "Sí", "Resolución judicial de adopción requerida."),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["Tipo de Licencia", "Duración", "Remunerada", "Requisitos"]):
        t.rows[0].cells[i].text = h
    for row_data in licencias:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    add_heading(doc, "3. Permisos", 2)
    add_paragraph(doc, "Los permisos personales permiten al colaborador ausentarse hasta 4 horas continuas, un máximo de 2 veces al mes. Los permisos personales no son remunerados y deben compensarse con horas adicionales en la misma semana, coordinadas con el jefe directo. Se solicitan por correo electrónico al jefe directo con un mínimo de 24 horas de anticipación, salvo emergencias.")

    add_heading(doc, "4. Contacto para Consultas", 2)
    add_paragraph(doc, "Gerencia de Recursos Humanos\nSilvia Torres Vega — s.torres@nexussoluciones.pe — Ext. 104\nLucía Romero Díaz (Selección) — l.romero@nexussoluciones.pe — Ext. 201\nJorge Palacios Ríos (Compensaciones) — j.palacios@nexussoluciones.pe — Ext. 202")

    doc.save(f"{OUTPUT_DIR}/D03_Politica_Vacaciones_Licencias.docx")
    print("D03 OK")

# ─────────────────────────────────────────────
# D04 — Política de Trabajo Remoto e Híbrido
# ─────────────────────────────────────────────
def gen_D04():
    doc = Document()
    header_footer(doc, "Política de Trabajo Remoto e Híbrido")
    add_heading(doc, "Política de Trabajo Remoto e Híbrido", 1)
    doc.add_paragraph("Código: RRHH-POL-003 | Versión: 2.1 | Fecha: Enero 2026 | Propietario: Gerencia de RR.HH. y CTO")

    add_heading(doc, "1. Modalidad Vigente", 2)
    add_paragraph(doc, "Nexus Soluciones S.A.C. opera bajo modalidad híbrida. Cada colaborador puede trabajar de forma remota hasta 2 días por semana, previa coordinación con su jefe directo.")

    add_heading(doc, "2. Elegibilidad", 2)
    add_paragraph(doc, "Para acceder al trabajo remoto el colaborador debe:")
    req = [
        "Haber completado los primeros 3 meses de permanencia en la empresa.",
        "Tener un desempeño evaluado como 'Cumple expectativas' o superior en su última evaluación.",
        "Contar con la aprobación de su jefe directo.",
        "Los roles que requieren presencia física permanente (soporte en sitio, acceso físico a equipos de infraestructura) no aplican a esta política.",
    ]
    for r in req:
        doc.add_paragraph(f"• {r}", style="List Bullet")

    add_heading(doc, "3. Requisitos Técnicos del Entorno Remoto", 2)
    tecn = [
        "Conexión a internet con velocidad mínima de 25 Mbps (subida y bajada).",
        "Espacio de trabajo privado, silencioso y adecuado para videollamadas.",
        "VPN corporativa activa: Cisco AnyConnect. Solicitar acceso a helpdesk@nexussoluciones.pe (Ext. 300).",
        "Disponibilidad completa en el horario laboral acordado (Turno A, B o C según corresponda).",
        "Equipo corporativo asignado por TI (no se permite uso de equipos personales para trabajo con datos de clientes).",
    ]
    for t_item in tecn:
        doc.add_paragraph(f"• {t_item}", style="List Bullet")

    add_heading(doc, "4. Equipamiento para el Hogar", 2)
    add_paragraph(doc, "La empresa provee un monitor externo adicional previa solicitud formal al área de TI a través del formulario en intranet.nexussoluciones.pe/solicitudes-ti. La solicitud requiere aprobación del jefe directo y disponibilidad de inventario.")

    add_heading(doc, "5. Proceso de Activación", 2)
    pasos = [
        "El colaborador acuerda los días de trabajo remoto con su jefe directo.",
        "El jefe directo registra el acuerdo en intranet.nexussoluciones.pe/trabajo-remoto.",
        "TI verifica que el colaborador tenga VPN activa y equipo corporativo en buen estado.",
        "RR.HH. valida el cumplimiento del período mínimo de permanencia.",
        "El acuerdo puede revisarse semestralmente o ante cambio de proyecto.",
    ]
    for i, p in enumerate(pasos, 1):
        doc.add_paragraph(f"{i}. {p}")

    add_heading(doc, "6. Obligaciones en Modalidad Remota", 2)
    oblig = [
        "Mantener la VPN activa durante toda la jornada laboral al usar sistemas internos.",
        "Responder mensajes en Teams y correo en máximo 30 minutos durante horario laboral.",
        "Asistir a reuniones con cámara encendida cuando sea solicitado.",
        "No trabajar desde espacios públicos (cafeterías, aeropuertos) con información confidencial o datos de clientes sin aprobación del CISO.",
        "Reportar cualquier incidente de seguridad inmediatamente a helpdesk@nexussoluciones.pe.",
    ]
    for o in oblig:
        doc.add_paragraph(f"• {o}", style="List Bullet")

    doc.save(f"{OUTPUT_DIR}/D04_Politica_Trabajo_Remoto.docx")
    print("D04 OK")

# ─────────────────────────────────────────────
# D05 — Evaluación de Desempeño y Compensaciones
# ─────────────────────────────────────────────
def gen_D05():
    doc = Document()
    header_footer(doc, "Evaluación de Desempeño y Compensaciones")
    add_heading(doc, "Reglamento de Evaluación de Desempeño y Compensaciones", 1)
    doc.add_paragraph("Código: RRHH-POL-004 | Versión: 3.2 | Fecha: Enero 2026 | Propietario: Gerencia de RR.HH.")

    add_heading(doc, "1. Ciclos de Evaluación", 2)
    add_paragraph(doc, "Nexus Soluciones realiza evaluaciones de desempeño dos veces al año bajo metodología 180 grados (autoevaluación + evaluación del jefe directo):")
    ciclos = [
        ("Evaluación Mid-Year", "Junio – Julio", "Agosto", "Revisión de objetivos del semestre, ajuste de plan de desarrollo"),
        ("Evaluación Year-End", "Noviembre – Diciembre", "Enero del año siguiente", "Evaluación anual, definición de incrementos salariales y bonos"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["Ciclo", "Período de Evaluación", "Resultados", "Propósito"]):
        t.rows[0].cells[i].text = h
    for row_data in ciclos:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    add_heading(doc, "2. Dimensiones Evaluadas", 2)
    dims = [
        ("Cumplimiento de objetivos", "Logro de los OKRs u objetivos acordados al inicio del período."),
        ("Competencias técnicas", "Dominio de las herramientas y conocimientos requeridos para el puesto."),
        ("Trabajo en equipo", "Colaboración, comunicación y contribución al clima del equipo."),
        ("Iniciativa e innovación", "Propuestas de mejora, resolución proactiva de problemas."),
        ("Comunicación", "Claridad, oportunidad y efectividad en la comunicación interna y con clientes."),
    ]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Dimensión"
    t2.rows[0].cells[1].text = "Descripción"
    for nombre, desc in dims:
        row = t2.add_row()
        row.cells[0].text = nombre
        row.cells[1].text = desc

    add_heading(doc, "3. Escala de Calificación", 2)
    escala = [
        ("1 — Necesita mejorar", "El colaborador no alcanza los estándares mínimos esperados en la dimensión."),
        ("2 — En desarrollo", "El colaborador está en proceso de alcanzar el nivel esperado, con apoyo."),
        ("3 — Cumple expectativas", "El colaborador cumple con lo esperado para su nivel y rol."),
        ("4 — Supera expectativas", "El colaborador excede consistentemente lo esperado."),
        ("5 — Excepcional", "El colaborador es un referente del área y genera impacto visible en la organización."),
    ]
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Puntaje"
    t3.rows[0].cells[1].text = "Descripción"
    for puntaje, desc in escala:
        row = t3.add_row()
        row.cells[0].text = puntaje
        row.cells[1].text = desc

    add_heading(doc, "4. Plataforma de Evaluación", 2)
    add_paragraph(doc, "Las evaluaciones se realizan en: intranet.nexussoluciones.pe/evaluacion. El sistema envía notificaciones automáticas al inicio de cada ciclo. Las evaluaciones incompletas al cierre del plazo se registran como 'No completada' y afectan el proceso de revisión salarial.")

    add_heading(doc, "5. Compensaciones y Beneficios Asociados", 2)
    add_paragraph(doc, "Revisión Salarial Anual:", bold=True)
    add_paragraph(doc, "Se realiza en enero de cada año, vinculada directamente al resultado de la Evaluación Year-End. Los incrementos se calculan sobre el resultado promedio de las cinco dimensiones.")
    add_paragraph(doc, "Bono de Desempeño Anual:", bold=True)
    add_paragraph(doc, "Los colaboradores con calificación promedio de 3.5 o superior son elegibles para el bono anual, equivalente a hasta 1.5 remuneraciones mensuales, proporcional al resultado obtenido.")
    add_paragraph(doc, "Beneficios de Ley:", bold=True)
    beneficios_ley = [
        "Gratificaciones: julio y diciembre (1 sueldo mensual cada una)",
        "CTS: depositada en mayo y noviembre",
        "ESSALUD: 9% a cargo de la empresa",
        "AFP u ONP: a elección del colaborador al ingreso",
    ]
    for b in beneficios_ley:
        doc.add_paragraph(f"• {b}", style="List Bullet")

    add_heading(doc, "6. Promociones", 2)
    add_paragraph(doc, "Las promociones de cargo están vinculadas a los resultados acumulados de dos evaluaciones consecutivas con calificación promedio mínima de 4.0, disponibilidad de la posición y aprobación de la Gerencia General. El proceso es gestionado por RR.HH. en coordinación con el Gerente del área.")

    doc.save(f"{OUTPUT_DIR}/D05_Evaluacion_Desempeno.docx")
    print("D05 OK")

# Run D01-D05
gen_D01()
gen_D02()
gen_D03()
gen_D04()
gen_D05()
print("Batch 1 (D01-D05) completado.")
